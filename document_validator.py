"""
Валидатор документов ВКР для проверки соблюдения требований форматирования
"""

from docx import Document
from docx.shared import Cm, Pt
from typing import Dict, Any, List, Tuple
from pathlib import Path
from dataclasses import dataclass, field

from style_based_classifier import StyleBasedClassifier
from document_state import logger


@dataclass
class ValidationIssue:
    """Класс для описания проблемы валидации"""
    type: str  # "error", "warning", "info"
    category: str  # "margins", "fonts", "headings", "lists", etc.
    description: str
    location: str  # где найдена проблема
    expected: str  # что ожидалось
    actual: str  # что найдено
    suggestion: str = ""  # рекомендация по исправлению


@dataclass
class ValidationReport:
    """Отчет о валидации документа"""
    total_issues: int = 0
    errors: int = 0
    warnings: int = 0
    info: int = 0
    issues: List[ValidationIssue] = field(default_factory=list)
    statistics: Dict[str, Any] = field(default_factory=dict)
    
    def add_issue(self, issue: ValidationIssue):
        """Добавляет проблему в отчет"""
        self.issues.append(issue)
        self.total_issues += 1
        
        if issue.type == "error":
            self.errors += 1
        elif issue.type == "warning":
            self.warnings += 1
        elif issue.type == "info":
            self.info += 1
    
    def get_summary(self) -> Dict[str, Any]:
        """Возвращает краткую сводку"""
        return {
            "total_issues": self.total_issues,
            "errors": self.errors,
            "warnings": self.warnings,
            "info": self.info,
            "score": self._calculate_score()
        }
    
    def _calculate_score(self) -> int:
        """Рассчитывает оценку соответствия (0-100)"""
        if self.total_issues == 0:
            return 100
        
        # Веса для разных типов проблем
        error_weight = 3
        warning_weight = 2
        info_weight = 1
        
        total_weight = (self.errors * error_weight + 
                       self.warnings * warning_weight + 
                       self.info * info_weight)
        
        # Базовая оценка зависит от количества проблем
        max_possible_weight = 50  # Условная максимальная сумма весов
        score = max(0, 100 - (total_weight * 100) // max_possible_weight)
        
        return min(100, score)


class DocumentValidator:
    """Валидатор документов ВКР"""

    def __init__(self, requirements: Dict[str, Any]):
        self.requirements = requirements
        self.classifier = StyleBasedClassifier(requirements, strict_style_mode=False)
        
    def validate_document(self, input_path: str) -> ValidationReport:
        """Валидирует документ на соответствие требованиям"""
        report = ValidationReport()
        
        try:
            logger.info(f"🔍 Начинаем валидацию: {input_path}")
            
            # Проверяем входной файл
            input_file = Path(input_path)
            if not input_file.exists():
                report.add_issue(ValidationIssue(
                    type="error",
                    category="file",
                    description="Файл не существует",
                    location=input_path,
                    expected="Существующий файл",
                    actual="Файл не найден"
                ))
                return report
            
            # Загружаем документ
            doc = Document(input_path)
            logger.info(f"📖 Документ загружен, параграфов: {len(doc.paragraphs)}")
            
            # Выполняем различные проверки
            self._validate_global_settings(doc, report)
            self._validate_paragraphs(doc, report)
            self._validate_tables(doc, report)
            self._validate_document_structure(doc, report)
            
            # Собираем статистику
            report.statistics = self._collect_statistics(doc)
            
            logger.info(f"✅ Валидация завершена. Найдено проблем: {report.total_issues}")
            
        except Exception as e:
            logger.error(f"Ошибка валидации: {e}")
            # Устанавливаем статистику по умолчанию если она еще не была собрана
            if not report.statistics:
                report.statistics = {
                    "total_paragraphs": 0,
                    "total_tables": 0,
                    "total_sections": 0,
                    "heading_counts": {"h1": 0, "h2": 0, "h3": 0, "h4": 0},
                    "list_items": 0,
                    "regular_paragraphs": 0,
                    "empty_paragraphs": 0
                }
            report.add_issue(ValidationIssue(
                type="error",
                category="system",
                description=f"Критическая ошибка валидации: {str(e)}",
                location="Системная ошибка",
                expected="Успешная валидация",
                actual=str(e)
            ))
        
        return report
    
    def _validate_global_settings(self, doc: Document, report: ValidationReport):
        """Проверяет глобальные настройки документа"""
        logger.info("🔍 Проверяем глобальные настройки...")
        
        base_config = self.requirements["base_formatting"]
        expected_margins = base_config["margins_cm"]
        
        for i, section in enumerate(doc.sections):
            # Проверяем поля с более точными измерениями
            actual_margins = {
                "top": round(section.top_margin.cm, 2),
                "bottom": round(section.bottom_margin.cm, 2),
                "left": round(section.left_margin.cm, 2),
                "right": round(section.right_margin.cm, 2)
            }
            
            for margin_name, expected_value in expected_margins.items():
                actual_value = actual_margins[margin_name]
                tolerance = 0.1  # Допуск в см
                
                if abs(actual_value - expected_value) > tolerance:
                    severity = "error" if abs(actual_value - expected_value) > 0.5 else "warning"
                    report.add_issue(ValidationIssue(
                        type=severity,
                        category="margins",
                        description=f"Неправильное {margin_name} поле документа",
                        location=f"Секция {i+1} (страницы документа)",
                        expected=f"{margin_name} поле: {expected_value} см",
                        actual=f"{margin_name} поле: {actual_value} см (отклонение: {abs(actual_value - expected_value):.2f} см)",
                        suggestion=f"Перейдите в Макет → Поля → Настраиваемые поля и установите {margin_name} поле = {expected_value} см"
                    ))
            
            # Проверяем ориентацию страницы
            if section.orientation != 0:  # 0 = portrait, 1 = landscape
                report.add_issue(ValidationIssue(
                    type="warning",
                    category="margins",
                    description="Неправильная ориентация страницы",
                    location=f"Секция {i+1}",
                    expected="Книжная ориентация",
                    actual="Альбомная ориентация",
                    suggestion="Измените ориентацию на книжную через Макет → Ориентация → Книжная"
                ))
            
            # Проверяем размер страницы
            page_width = section.page_width.cm
            page_height = section.page_height.cm
            # A4: 21.0 x 29.7 см
            if abs(page_width - 21.0) > 0.5 or abs(page_height - 29.7) > 0.5:
                report.add_issue(ValidationIssue(
                    type="warning",
                    category="margins",
                    description="Неправильный размер страницы",
                    location=f"Секция {i+1}",
                    expected="Формат A4 (21.0 × 29.7 см)",
                    actual=f"Текущий размер: {page_width:.1f} × {page_height:.1f} см",
                    suggestion="Установите формат A4 через Макет → Размер → A4"
                ))
        
        # Проверяем стили документа
        self._validate_default_styles(doc, report, base_config)
        
        # Проверяем нумерацию страниц
        self._check_page_numbering(doc, report)
    
    def _validate_default_styles(self, doc: Document, report: ValidationReport, base_config: Dict[str, Any]):
        """Проверяет стили документа по умолчанию"""
        try:
            normal_style = doc.styles['Normal']
            font = normal_style.font
            
            # Проверяем шрифт
            expected_font = base_config["font_name"]
            if font.name != expected_font:
                report.add_issue(ValidationIssue(
                    type="error",
                    category="fonts",
                    description="Неправильный шрифт в стиле Normal",
                    location="Стиль Normal",
                    expected=expected_font,
                    actual=font.name or "Не установлен",
                    suggestion=f"Установите шрифт '{expected_font}' для стиля Normal"
                ))
            
            # Проверяем размер шрифта
            expected_size = base_config["font_size"]
            actual_size = font.size.pt if font.size else None
            
            if actual_size != expected_size:
                report.add_issue(ValidationIssue(
                    type="error",
                    category="fonts",
                    description="Неправильный размер шрифта в стиле Normal",
                    location="Стиль Normal",
                    expected=f"{expected_size} пт",
                    actual=f"{actual_size} пт" if actual_size else "Не установлен",
                    suggestion=f"Установите размер шрифта {expected_size} пт для стиля Normal"
                ))
                
        except Exception as e:
            report.add_issue(ValidationIssue(
                type="warning",
                category="fonts",
                description=f"Не удалось проверить стиль Normal: {str(e)}",
                location="Стили документа",
                expected="Доступный стиль Normal",
                actual="Ошибка доступа к стилю"
            ))
    
    def _validate_paragraphs(self, doc: Document, report: ValidationReport):
        """Проверяет форматирование параграфов"""
        logger.info("🔍 Проверяем параграфы...")
        
        heading_counts = {"h1": 0, "h2": 0, "h3": 0, "h4": 0}
        list_items = 0
        regular_paragraphs = 0
        
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue
                
            # Классифицируем параграф
            paragraph_type = self.classifier.classify_paragraph_by_style(paragraph, text)
            
            # Получаем информацию о стиле параграфа
            style_name = self._get_paragraph_style_name(paragraph)
            text_preview = text[:50] + "..." if len(text) > 50 else text
            
            if paragraph_type in heading_counts:
                heading_counts[paragraph_type] += 1
                self._validate_heading(paragraph, paragraph_type, i+1, report, text_preview, style_name)
            elif paragraph_type == "list":
                list_items += 1
                self._validate_list_item(paragraph, i+1, report, text_preview)
            elif paragraph_type == "regular":
                regular_paragraphs += 1
                self._validate_regular_paragraph(paragraph, i+1, report, text_preview)
            elif paragraph_type == "skip":
                # Дополнительная информация о пропущенных параграфах
                logger.debug(f"   🔄 Пропущен параграф {i+1}: {text_preview}")
            else:
                # Неопознанные типы параграфов
                report.add_issue(ValidationIssue(
                    type="info",
                    category="structure",
                    description=f"Неопознанный тип параграфа",
                    location=f"Параграф {i+1}: \"{text_preview}\"",
                    expected="Опознанный тип (заголовок, список или обычный текст)",
                    actual=f"Тип: {paragraph_type}, Стиль: {style_name}",
                    suggestion="Проверьте форматирование или стиль параграфа"
                ))
        
        # Проверяем структуру документа
        if heading_counts["h1"] == 0:
            report.add_issue(ValidationIssue(
                type="warning",
                category="structure",
                description="В документе нет заголовков первого уровня",
                location="Весь документ",
                expected="Минимум 1 заголовок H1 (главы документа)",
                actual="0 заголовков H1",
                suggestion="Добавьте заголовки глав (H1) для основных разделов: Введение, Глава 1, Глава 2, Заключение"
            ))
            
        # Проверяем баланс заголовков
        total_headings = sum(heading_counts.values())
        if total_headings > 0:
            h1_ratio = heading_counts["h1"] / total_headings
            if h1_ratio > 0.5:
                report.add_issue(ValidationIssue(
                    type="info",
                    category="structure",
                    description="Слишком много заголовков H1 относительно других уровней",
                    location=f"H1: {heading_counts['h1']}, H2: {heading_counts['h2']}, H3: {heading_counts['h3']}",
                    expected="H1 составляют 20-40% от всех заголовков",
                    actual=f"H1 составляют {h1_ratio*100:.1f}% от всех заголовков",
                    suggestion="Рассмотрите возможность объединения некоторых глав или использования заголовков H2-H3"
                ))

    def _get_paragraph_style_name(self, paragraph) -> str:
        """Получает название стиля параграфа с улучшенной обработкой ошибок"""
        try:
            if paragraph.style and paragraph.style.name:
                return paragraph.style.name
            else:
                return "Normal (по умолчанию)"
        except Exception as e:
            return f"Ошибка стиля: {str(e)}"

    def _validate_heading(self, paragraph, heading_type: str, paragraph_num: int, report: ValidationReport, text_preview: str, style_name: str):
        """Проверяет форматирование заголовка с улучшенной детализацией"""
        requirements = self.requirements[f"{heading_type}_formatting"]
        
        # Проверяем каждый run в параграфе
        for run_index, run in enumerate(paragraph.runs):
            run_text = run.text.strip()
            if not run_text:  # Пропускаем пустые runs
                continue
                
            # Проверяем шрифт
            actual_font = run.font.name or "Не установлен"
            expected_font = requirements["font_name"]
            if actual_font != expected_font:
                report.add_issue(ValidationIssue(
                    type="error",
                    category="headings",
                    description=f"Неправильный шрифт в заголовке {heading_type.upper()}",
                    location=f"Параграф {paragraph_num}, фрагмент {run_index+1}: \"{text_preview}\"",
                    expected=f"Шрифт: {expected_font}",
                    actual=f"Шрифт: {actual_font}, Стиль параграфа: {style_name}",
                    suggestion=f"Выделите заголовок и установите шрифт '{expected_font}' через меню Главная → Шрифт"
                ))
            
            # Проверяем размер шрифта
            actual_size = run.font.size.pt if run.font.size else None
            expected_size = requirements["font_size"]
            if actual_size and actual_size != expected_size:
                report.add_issue(ValidationIssue(
                    type="error",
                    category="headings",
                    description=f"Неправильный размер шрифта в заголовке {heading_type.upper()}",
                    location=f"Параграф {paragraph_num}: \"{text_preview}\"",
                    expected=f"Размер: {expected_size} пт",
                    actual=f"Размер: {actual_size} пт, Стиль: {style_name}",
                    suggestion=f"Выделите заголовок и установите размер {expected_size} пт через меню Главная → Размер шрифта"
                ))
            elif not actual_size:
                report.add_issue(ValidationIssue(
                    type="warning",
                    category="headings",
                    description=f"Размер шрифта не установлен для заголовка {heading_type.upper()}",
                    location=f"Параграф {paragraph_num}: \"{text_preview}\"",
                    expected=f"Размер: {expected_size} пт",
                    actual="Размер не определен",
                    suggestion=f"Установите размер шрифта {expected_size} пт"
                ))
            
            # Проверяем жирность
            is_bold = run.font.bold
            should_be_bold = requirements.get("bold", False)
            if should_be_bold and not is_bold:
                report.add_issue(ValidationIssue(
                    type="error",
                    category="headings",
                    description=f"Заголовок {heading_type.upper()} должен быть жирным",
                    location=f"Параграф {paragraph_num}: \"{text_preview}\"",
                    expected="Жирный шрифт (Bold)",
                    actual=f"Обычный шрифт, Стиль: {style_name}",
                    suggestion="Выделите заголовок и нажмите Ctrl+B или кнопку 'Ж' на панели инструментов"
                ))
            
            # Проверяем курсив (если не должен быть)
            is_italic = run.font.italic
            if is_italic and not requirements.get("italic", False):
                report.add_issue(ValidationIssue(
                    type="warning",
                    category="headings",
                    description=f"Заголовок {heading_type.upper()} не должен быть курсивным",
                    location=f"Параграф {paragraph_num}: \"{text_preview}\"",
                    expected="Обычный шрифт (без курсива)",
                    actual="Курсивный шрифт",
                    suggestion="Снимите курсив с заголовка (Ctrl+I)"
                ))
            
            # Проверяем подчеркивание (обычно не нужно)
            is_underlined = run.font.underline
            if is_underlined:
                report.add_issue(ValidationIssue(
                    type="info",
                    category="headings",
                    description=f"Заголовок {heading_type.upper()} подчеркнут",
                    location=f"Параграф {paragraph_num}: \"{text_preview}\"",
                    expected="Заголовок без подчеркивания",
                    actual="Подчеркнутый текст",
                    suggestion="Обычно заголовки не подчеркивают. Снимите подчеркивание (Ctrl+U)"
                ))
        
        # Проверяем выравнивание
        expected_alignment = requirements.get("alignment", "center")
        self._check_alignment(paragraph, expected_alignment, f"заголовка {heading_type.upper()}", 
                             paragraph_num, report, text_preview)
        
        # Проверяем межстрочный интервал
        line_spacing = paragraph.paragraph_format.line_spacing
        if line_spacing and line_spacing != 1.0:
            report.add_issue(ValidationIssue(
                type="warning",
                category="headings",
                description=f"Неправильный межстрочный интервал в заголовке {heading_type.upper()}",
                location=f"Параграф {paragraph_num}: \"{text_preview}\"",
                expected="Одинарный интервал (1.0)",
                actual=f"Интервал: {line_spacing}",
                suggestion="Установите одинарный межстрочный интервал для заголовков"
            ))

    def _validate_list_item(self, paragraph, paragraph_num: int, report: ValidationReport, text_preview: str):
        """Проверяет форматирование элемента списка с детализацией"""
        list_config = self.requirements["lists"]["bullet_lists"]
        
        # Проверяем отступы
        expected_indent = Cm(list_config["indent_cm"])
        actual_indent = paragraph.paragraph_format.left_indent
        actual_cm = actual_indent.cm if actual_indent else 0
        
        if abs(actual_cm - list_config["indent_cm"]) > 0.1:  # Допуск 1мм
            report.add_issue(ValidationIssue(
                type="warning",
                category="lists",
                description="Неправильный отступ в элементе списка",
                location=f"Параграф {paragraph_num}: \"{text_preview}\"",
                expected=f"Отступ: {list_config['indent_cm']} см",
                actual=f"Отступ: {actual_cm:.2f} см",
                suggestion=f"Установите отступ {list_config['indent_cm']} см через Главная → Увеличить отступ или Формат → Абзац"
            ))
        
        # Проверяем висячий отступ
        hanging_indent = paragraph.paragraph_format.first_line_indent
        if hanging_indent and hanging_indent.cm < -0.1:  # Есть висячий отступ
            report.add_issue(ValidationIssue(
                type="info",
                category="lists",
                description="Обнаружен висячий отступ в элементе списка",
                location=f"Параграф {paragraph_num}: \"{text_preview}\"",
                expected="Обычный отступ без висячего",
                actual=f"Висячий отступ: {hanging_indent.cm:.2f} см",
                suggestion="Используйте автоматическое форматирование списков вместо ручных отступов"
            ))
        
        # Проверяем межстрочный интервал
        line_spacing = paragraph.paragraph_format.line_spacing
        expected_spacing = 1.0
        if line_spacing and abs(line_spacing - expected_spacing) > 0.1:
            report.add_issue(ValidationIssue(
                type="warning",
                category="lists",
                description="Неправильный межстрочный интервал в списке",
                location=f"Параграф {paragraph_num}: \"{text_preview}\"",
                expected=f"Интервал: {expected_spacing}",
                actual=f"Интервал: {line_spacing}",
                suggestion="Установите одинарный межстрочный интервал для элементов списка"
            ))

    def _validate_regular_paragraph(self, paragraph, paragraph_num: int, report: ValidationReport, text_preview: str):
        """Проверяет форматирование обычного параграфа с детализацией"""
        base_config = self.requirements["base_formatting"]
        
        # Проверяем красную строку
        expected_indent = Cm(base_config["first_line_indent_cm"])
        actual_indent = paragraph.paragraph_format.first_line_indent
        actual_cm = actual_indent.cm if actual_indent else 0
        
        if abs(actual_cm - base_config["first_line_indent_cm"]) > 0.1:  # Допуск 1мм
            report.add_issue(ValidationIssue(
                type="warning",
                category="paragraphs",
                description="Неправильная красная строка",
                location=f"Параграф {paragraph_num}: \"{text_preview}\"",
                expected=f"Красная строка: {base_config['first_line_indent_cm']} см",
                actual=f"Красная строка: {actual_cm:.2f} см",
                suggestion=f"Установите красную строку {base_config['first_line_indent_cm']} см через Формат → Абзац → Первая строка: Отступ"
            ))
        
        # Проверяем выравнивание
        self._check_alignment(paragraph, "justify", "обычного параграфа", paragraph_num, report, text_preview)
        
        # Проверяем межстрочный интервал
        line_spacing = paragraph.paragraph_format.line_spacing
        expected_spacing = base_config.get("line_spacing", 1.5)
        if line_spacing and abs(line_spacing - expected_spacing) > 0.1:
            report.add_issue(ValidationIssue(
                type="warning",
                category="paragraphs",
                description="Неправильный межстрочный интервал",
                location=f"Параграф {paragraph_num}: \"{text_preview}\"",
                expected=f"Интервал: {expected_spacing}",
                actual=f"Интервал: {line_spacing}",
                suggestion=f"Установите междустрочный интервал {expected_spacing} через Формат → Абзац → Междустрочный интервал"
            ))
        
        # Проверяем шрифт в обычном тексте
        for run in paragraph.runs:
            if run.text.strip():  # Если есть текст
                # Проверяем шрифт
                actual_font = run.font.name
                expected_font = base_config["font_name"]
                if actual_font and actual_font != expected_font:
                    report.add_issue(ValidationIssue(
                        type="error",
                        category="fonts",
                        description="Неправильный шрифт в обычном тексте",
                        location=f"Параграф {paragraph_num}: \"{text_preview}\"",
                        expected=f"Шрифт: {expected_font}",
                        actual=f"Шрифт: {actual_font}",
                        suggestion=f"Выделите текст и установите шрифт '{expected_font}'"
                    ))
                
                # Проверяем размер
                actual_size = run.font.size.pt if run.font.size else None
                expected_size = base_config["font_size"]
                if actual_size and actual_size != expected_size:
                    report.add_issue(ValidationIssue(
                        type="error",
                        category="fonts",
                        description="Неправильный размер шрифта в обычном тексте",
                        location=f"Параграф {paragraph_num}: \"{text_preview}\"",
                        expected=f"Размер: {expected_size} пт",
                        actual=f"Размер: {actual_size} пт",
                        suggestion=f"Установите размер шрифта {expected_size} пт"
                    ))
                break  # Проверяем только первый run с текстом

    def _check_alignment(self, paragraph, expected_alignment: str, element_type: str, 
                        paragraph_num: int, report: ValidationReport, text_preview: str = ""):
        """Проверяет выравнивание параграфа с улучшенной детализацией"""
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        alignment_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        
        alignment_names = {
            WD_ALIGN_PARAGRAPH.LEFT: "по левому краю",
            WD_ALIGN_PARAGRAPH.CENTER: "по центру", 
            WD_ALIGN_PARAGRAPH.RIGHT: "по правому краю",
            WD_ALIGN_PARAGRAPH.JUSTIFY: "по ширине",
            None: "не установлено"
        }
        
        expected_enum = alignment_map.get(expected_alignment)
        actual_enum = paragraph.paragraph_format.alignment
        
        if actual_enum != expected_enum:
            actual_name = alignment_names.get(actual_enum, "неизвестное")
            expected_name = alignment_names.get(expected_enum, expected_alignment)
            
            location = f"Параграф {paragraph_num}"
            if text_preview:
                location += f": \"{text_preview}\""
            
            report.add_issue(ValidationIssue(
                type="warning",
                category="alignment",
                description=f"Неправильное выравнивание {element_type}",
                location=location,
                expected=f"Выравнивание: {expected_name}",
                actual=f"Выравнивание: {actual_name}",
                suggestion=f"Выделите текст и установите выравнивание '{expected_name}' через Главная → Выравнивание или Ctrl+E/L/R/J"
            ))
    
    def _validate_tables(self, doc: Document, report: ValidationReport):
        """Проверяет форматирование таблиц с детализацией"""
        logger.info("🔍 Проверяем таблицы...")
        
        if not doc.tables:
            report.add_issue(ValidationIssue(
                type="info",
                category="tables",
                description="В документе нет таблиц",
                location="Весь документ",
                expected="Возможно, таблицы нужны для представления данных",
                actual="Таблицы отсутствуют",
                suggestion="Если в работе есть табличные данные, оформите их в виде таблиц"
            ))
            return
        
        table_config = self.requirements.get("tables", {})
        
        for i, table in enumerate(doc.tables):
            table_location = f"Таблица {i+1}"
            
            # Проверяем наличие строк
            if not table.rows:
                report.add_issue(ValidationIssue(
                    type="warning",
                    category="tables",
                    description="Пустая таблица",
                    location=table_location,
                    expected="Таблица с данными (заголовки + строки данных)",
                    actual="Таблица без строк",
                    suggestion="Заполните таблицу данными или удалите её"
                ))
                continue
            
            # Анализируем содержимое таблицы
            total_cells = 0
            empty_cells = 0
            
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    total_cells += 1
                    cell_text = cell.text.strip()
                    
                    if not cell_text:
                        empty_cells += 1
                    
                    # Проверяем форматирование ячеек
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            # Проверяем выравнивание в ячейках
                            if paragraph.paragraph_format.alignment is None:
                                report.add_issue(ValidationIssue(
                                    type="info",
                                    category="tables",
                                    description="Не установлено выравнивание в ячейке таблицы",
                                    location=f"{table_location}, строка {row_idx+1}, столбец {cell_idx+1}",
                                    expected="Явно установленное выравнивание (обычно по центру для заголовков)",
                                    actual="Выравнивание не установлено",
                                    suggestion="Установите выравнивание для ячеек таблицы"
                                ))
                            
                            # Проверяем шрифт в ячейках
                            for run in paragraph.runs:
                                if run.text.strip() and run.font.name:
                                    expected_font = self.requirements["base_formatting"]["font_name"]
                                    if run.font.name != expected_font:
                                        report.add_issue(ValidationIssue(
                                            type="warning",
                                            category="tables",
                                            description="Неправильный шрифт в таблице",
                                            location=f"{table_location}, строка {row_idx+1}, столбец {cell_idx+1}",
                                            expected=f"Шрифт: {expected_font}",
                                            actual=f"Шрифт: {run.font.name}",
                                            suggestion=f"Установите шрифт '{expected_font}' для всей таблицы"
                                        ))
                                break
                            break
            
            # Проверяем заполненность таблицы
            if total_cells > 0:
                empty_ratio = empty_cells / total_cells
                if empty_ratio > 0.3:  # Более 30% пустых ячеек
                    report.add_issue(ValidationIssue(
                        type="warning",
                        category="tables",
                        description="Много пустых ячеек в таблице",
                        location=f"{table_location} ({len(table.rows)} строк, {len(table.columns)} столбцов)",
                        expected="Заполненная таблица (менее 30% пустых ячеек)",
                        actual=f"Пустых ячеек: {empty_cells} из {total_cells} ({empty_ratio*100:.1f}%)",
                        suggestion="Заполните пустые ячейки данными или удалите лишние строки/столбцы"
                    ))
            
            # Проверяем заголовки таблицы
            if len(table.rows) > 0:
                first_row = table.rows[0]
                has_headers = False
                
                for cell in first_row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if run.font.bold:
                                has_headers = True
                                break
                        if has_headers:
                            break
                    if has_headers:
                        break
                
                if not has_headers and len(table.rows) > 1:
                    report.add_issue(ValidationIssue(
                        type="warning",
                        category="tables",
                        description="Первая строка таблицы не выглядит как заголовок",
                        location=table_location,
                        expected="Заголовки таблицы выделены жирным шрифтом",
                        actual="Первая строка не выделена",
                        suggestion="Выделите заголовки таблицы жирным шрифтом или примените стиль 'Заголовок таблицы'"
                    ))
            
            # Проверяем подпись к таблице
            self._check_table_caption(doc, i, report)

    def _check_table_caption(self, doc: Document, table_index: int, report: ValidationReport):
        """Проверяет наличие подписи к таблице"""
        try:
            # Ищем параграфы до и после таблицы, которые могут быть подписями
            found_caption = False
            
            # Простая эвристика: ищем текст с "Таблица" рядом с таблицей
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip().upper()
                if "ТАБЛИЦА" in text and str(table_index + 1) in text:
                    found_caption = True
                    break
            
            if not found_caption:
                report.add_issue(ValidationIssue(
                    type="info",
                    category="tables",
                    description=f"Не найдена подпись к таблице {table_index + 1}",
                    location=f"Таблица {table_index + 1}",
                    expected="Подпись вида 'Таблица 1 - Название таблицы'",
                    actual="Подпись не обнаружена",
                    suggestion="Добавьте подпись к таблице согласно ГОСТ (перед таблицей или после неё)"
                ))
        except Exception as e:
            logger.debug(f"Ошибка проверки подписи таблицы: {e}")

    def _validate_document_structure(self, doc: Document, report: ValidationReport):
        """Проверяет общую структуру документа"""
        logger.info("🔍 Проверяем структуру документа...")
        
        # Проверяем последовательность заголовков
        heading_levels = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
                
            paragraph_type = self.classifier.classify_paragraph_by_style(paragraph, text)
            if paragraph_type in ["h1", "h2", "h3", "h4"]:
                level = int(paragraph_type[1])
                heading_levels.append(level)
        
        # Проверяем корректность последовательности
        for i in range(1, len(heading_levels)):
            current_level = heading_levels[i]
            prev_level = heading_levels[i-1]
            
            # Проверяем, что уровень не увеличивается более чем на 1
            if current_level > prev_level + 1:
                report.add_issue(ValidationIssue(
                    type="warning",
                    category="structure",
                    description=f"Пропущен уровень заголовков: с H{prev_level} сразу на H{current_level}",
                    location=f"Заголовок №{i+1}",
                    expected=f"H{prev_level+1} или меньше",
                    actual=f"H{current_level}",
                    suggestion=f"Добавьте промежуточные заголовки или измените уровень на H{prev_level+1}"
                ))
    
    def _collect_statistics(self, doc: Document) -> Dict[str, Any]:
        """Собирает статистику документа"""
        stats = {
            "total_paragraphs": len(doc.paragraphs),
            "total_tables": len(doc.tables),
            "total_sections": len(doc.sections),
            "heading_counts": {"h1": 0, "h2": 0, "h3": 0, "h4": 0},
            "list_items": 0,
            "regular_paragraphs": 0,
            "empty_paragraphs": 0
        }
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                stats["empty_paragraphs"] += 1
                continue
                
            paragraph_type = self.classifier.classify_paragraph_by_style(paragraph, text)
            if paragraph_type in stats["heading_counts"]:
                stats["heading_counts"][paragraph_type] += 1
            elif paragraph_type == "list":
                stats["list_items"] += 1
            elif paragraph_type == "regular":
                stats["regular_paragraphs"] += 1
        
        return stats

    def _check_page_numbering(self, doc: Document, report: ValidationReport):
        """Проверяет настройки нумерации страниц"""
        try:
            has_page_numbers = False
            for section in doc.sections:
                # Проверяем колонтитулы на наличие номеров страниц
                if section.header or section.footer:
                    # Простая проверка - есть ли текст в колонтитулах
                    if (section.header and section.header.paragraphs and 
                        any(p.text.strip() for p in section.header.paragraphs)):
                        has_page_numbers = True
                        break
                    if (section.footer and section.footer.paragraphs and 
                        any(p.text.strip() for p in section.footer.paragraphs)):
                        has_page_numbers = True
                        break
            
            if not has_page_numbers:
                report.add_issue(ValidationIssue(
                    type="info",
                    category="structure",
                    description="Нумерация страниц не обнаружена",
                    location="Колонтитулы документа",
                    expected="Нумерация страниц в верхнем или нижнем колонтитуле",
                    actual="Колонтитулы пустые или отсутствуют",
                    suggestion="Добавьте нумерацию страниц через Вставка → Номер страницы"
                ))
        except Exception as e:
            logger.debug(f"Ошибка проверки нумерации страниц: {e}")


def validate_vkr_document(input_path: str, requirements: Dict[str, Any]) -> Tuple[bool, ValidationReport]:
    """
    Валидирует ВКР документ
    
    Returns:
        Tuple[bool, ValidationReport]: (успех, отчет о валидации)
    """
    try:
        validator = DocumentValidator(requirements)
        report = validator.validate_document(input_path)
        
        # Считаем валидацию успешной, если нет критических ошибок
        success = report.errors == 0
        
        return success, report
        
    except Exception as e:
        logger.error(f"Критическая ошибка валидации: {e}")
        error_report = ValidationReport()
        # Устанавливаем пустую статистику для избежания ошибок Pydantic
        error_report.statistics = {
            "total_paragraphs": 0,
            "total_tables": 0,
            "total_sections": 0,
            "heading_counts": {"h1": 0, "h2": 0, "h3": 0, "h4": 0},
            "list_items": 0,
            "regular_paragraphs": 0,
            "empty_paragraphs": 0
        }
        error_report.add_issue(ValidationIssue(
            type="error",
            category="system",
            description=f"Критическая ошибка: {str(e)}",
            location="Системная ошибка",
            expected="Успешная валидация",
            actual=str(e)
        ))
        return False, error_report 