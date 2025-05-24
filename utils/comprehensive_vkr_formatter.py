from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from openai import OpenAI
import json
import re
import logging
from typing import Dict, Any, List, Optional, Tuple
import traceback
from dotenv import load_dotenv
import os

# Load environment variables
load_dotenv()

# Настройка логирования
logging.basicConfig(level=logging.INFO,
                    format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)


class VKRAnalyzer:
    """Анализатор ВКР для извлечения требований по полной структуре"""

    def __init__(self):
        api_key = os.getenv('OPENAI_API_KEY')
        if not api_key:
            raise ValueError(
                "OPENAI_API_KEY not found in environment variables")
        self.client = OpenAI(api_key=api_key)

    def analyze_requirements(self, requirements_path: str) -> Dict[str, Any]:
        """Анализирует файл требований и возвращает полную JSON структуру"""

        try:
            # Читаем документ с требованиями
            doc = Document(requirements_path)
            requirements_text = "\n".join(
                [p.text for p in doc.paragraphs if p.text.strip()])

            logger.info("Начинаем комплексный анализ требований...")

            # Получаем структурированные требования от ИИ
            requirements_json = self._extract_comprehensive_requirements(
                requirements_text)

            logger.info("Комплексный анализ завершен")
            return requirements_json

        except Exception as e:
            logger.error(f"Ошибка анализа требований: {e}")
            return self._get_default_requirements()

    def _extract_comprehensive_requirements(self, requirements_text: str) -> Dict[str, Any]:
        """Извлекает все требования через ИИ согласно нашей структуре"""

        prompt = f"""Проанализируй требования к оформлению дипломной работы и извлеки ВСЕ параметры форматирования:

ТРЕБОВАНИЯ:
---
{requirements_text}
---

Извлеки требования по следующим категориям:
1. ГЛОБАЛЬНЫЕ НАСТРОЙКИ (поля, основной текст, нумерация страниц)
2. ЗАГОЛОВКИ H1 (главы - размер, выравнивание, разрывы страниц)
3. ЗАГОЛОВКИ H2 (подразделы - размер, отступы)
4. СПЕЦИАЛЬНЫЕ РАЗДЕЛЫ (реферат, введение, заключение, содержание, список литературы)
5. РИСУНКИ (подписи, нумерация, выравнивание)
6. ТАБЛИЦЫ (подписи, размер шрифта, выравнивание)
7. ФОРМУЛЫ (нумерация, выравнивание, объяснения)
8. СПИСКИ (маркеры, пунктуация, отступы)

Верни только JSON без комментариев."""

        # Схема для полного анализа
        function_schema = {
            "name": "comprehensive_vkr_requirements",
            "description": "Полное извлечение требований к оформлению ВКР",
            "parameters": {
                "type": "object",
                "properties": {
                    "global_formatting": {
                        "type": "object",
                        "properties": {
                            "page_settings": {
                                "type": "object",
                                "properties": {
                                    "margins_cm": {
                                        "type": "object",
                                        "properties": {
                                            "top": {"type": "number"},
                                            "bottom": {"type": "number"},
                                            "left": {"type": "number"},
                                            "right": {"type": "number"}
                                        }
                                    }
                                }
                            },
                            "base_text": {
                                "type": "object",
                                "properties": {
                                    "font_name": {"type": "string"},
                                    "font_size": {"type": "integer"},
                                    "line_spacing": {"type": "number"},
                                    "alignment": {"type": "string"},
                                    "paragraph_indent_cm": {"type": "number"}
                                }
                            }
                        }
                    },
                    "headings": {
                        "type": "object",
                        "properties": {
                            "h1": {
                                "type": "object",
                                "properties": {
                                    "font_name": {"type": "string"},
                                    "font_size": {"type": "integer"},
                                    "font_style": {
                                        "type": "object",
                                        "properties": {
                                            "bold": {"type": "boolean"},
                                            "uppercase": {"type": "boolean"}
                                        }
                                    },
                                    "alignment": {"type": "string"},
                                    "page_break": {
                                        "type": "object",
                                        "properties": {
                                            "before": {"type": "boolean"}
                                        }
                                    }
                                }
                            },
                            "h2": {
                                "type": "object",
                                "properties": {
                                    "font_name": {"type": "string"},
                                    "font_size": {"type": "integer"},
                                    "font_style": {
                                        "type": "object",
                                        "properties": {
                                            "bold": {"type": "boolean"}
                                        }
                                    },
                                    "alignment": {"type": "string"}
                                }
                            }
                        }
                    },
                    "lists": {
                        "type": "object",
                        "properties": {
                            "bullet_lists": {
                                "type": "object",
                                "properties": {
                                    "marker": {"type": "string"},
                                    "indent_cm": {"type": "number"},
                                    "punctuation": {
                                        "type": "object",
                                        "properties": {
                                            "item_ending": {"type": "string"},
                                            "last_item_ending": {"type": "string"}
                                        }
                                    }
                                }
                            }
                        }
                    },
                    "figures": {
                        "type": "object",
                        "properties": {
                            "caption": {
                                "type": "object",
                                "properties": {
                                    "position": {"type": "string"},
                                    "alignment": {"type": "string"}
                                }
                            }
                        }
                    },
                    "tables": {
                        "type": "object",
                        "properties": {
                            "caption": {
                                "type": "object",
                                "properties": {
                                    "position": {"type": "string"},
                                    "alignment": {"type": "string"}
                                }
                            },
                            "content": {
                                "type": "object",
                                "properties": {
                                    "font_size": {"type": "integer"}
                                }
                            }
                        }
                    }
                },
                "required": ["global_formatting", "headings"]
            }
        }

        try:
            response = self.client.chat.completions.create(
                model="gpt-3.5-turbo-1106",
                messages=[{"role": "user", "content": prompt}],
                tools=[{"type": "function", "function": function_schema}],
                tool_choice="required",
            )

            args = response.choices[0].message.tool_calls[0].function.arguments
            requirements = json.loads(args)

            # Дополняем недостающие поля значениями по умолчанию
            requirements = self._merge_with_defaults(requirements)

            logger.info("ИИ успешно извлек требования")
            return requirements

        except Exception as e:
            logger.error(f"Ошибка извлечения требований через ИИ: {e}")
            return self._get_default_requirements()

    def _merge_with_defaults(self, requirements: Dict[str, Any]) -> Dict[str, Any]:
        """Дополняет извлеченные требования значениями по умолчанию"""

        defaults = self._get_default_requirements()

        def deep_merge(default_dict, user_dict):
            result = default_dict.copy()
            for key, value in user_dict.items():
                if key in result and isinstance(result[key], dict) and isinstance(value, dict):
                    result[key] = deep_merge(result[key], value)
                else:
                    result[key] = value
            return result

        return deep_merge(defaults, requirements)

    def _get_default_requirements(self) -> Dict[str, Any]:
        """Возвращает требования по умолчанию согласно ГОСТ"""

        return {
            "document_analysis": {
                "skip_sections": {
                    "title_page": True,
                    "task_assignment": True,
                    "calendar_plan": True,
                    "definitions_abbreviations": True
                }
            },
            "global_formatting": {
                "page_settings": {
                    "margins_cm": {
                        "top": 2.0,
                        "bottom": 2.0,
                        "left": 3.0,
                        "right": 1.5
                    }
                },
                "base_text": {
                    "font_name": "Times New Roman",
                    "font_size": 14,
                    "line_spacing": 1.5,
                    "alignment": "justify",
                    "paragraph_indent_cm": 1.25
                }
            },
            "headings": {
                "h1": {
                    "font_name": "Times New Roman",
                    "font_size": 16,
                    "font_style": {
                        "bold": True,
                        "uppercase": True
                    },
                    "alignment": "center",
                    "page_break": {
                        "before": True
                    },
                    "spacing": {
                        "before_pt": 18,
                        "after_pt": 12
                    }
                },
                "h2": {
                    "font_name": "Times New Roman",
                    "font_size": 14,
                    "font_style": {
                        "bold": True
                    },
                    "alignment": "left",
                    "spacing": {
                        "before_pt": 12,
                        "after_pt": 6
                    }
                }
            },
            "lists": {
                "bullet_lists": {
                    "marker": "–",
                    "indent_cm": 1.25,
                    "punctuation": {
                        "item_ending": ";",
                        "last_item_ending": "."
                    },
                    "font": {
                        "name": "Times New Roman",
                        "size": 14,
                        "line_spacing": 1.5
                    },
                    "alignment": "justify"
                }
            },
            "figures": {
                "caption": {
                    "position": "below",
                    "alignment": "center",
                    "font_size": 14
                },
                "alignment": "center"
            },
            "tables": {
                "caption": {
                    "position": "above",
                    "alignment": "left",
                    "font_size": 14
                },
                "content": {
                    "font_size": 12,
                    "alignment": "center"
                }
            },
            "detection_rules": {
                "h1_patterns": [
                    r"^\d+\.\s*[А-ЯЁ\s]+$",
                    r"^ГЛАВА\s+\d+",
                    r"^[IVX]+\.\s*[А-ЯЁ\s]+$",
                    r"^(ВВЕДЕНИЕ|ЗАКЛЮЧЕНИЕ|РЕФЕРАТ)$"
                ],
                "h2_patterns": [
                    r"^\d+\.\d+\.?\s+[А-Яа-яёЁ]",
                    r"^\d+\.\d+\s+[А-ЯЁ\s]+$"
                ],
                "list_patterns": [
                    r"^\s*[-–—]\s+",
                    r"^\s*\d+\)\s+",
                    r"^\s*[а-я]\)\s+"
                ],
                "skip_sections": {
                    "title_page": ["ДИПЛОМНАЯ РАБОТА", "ВЫПУСКНАЯ КВАЛИФИКАЦИОННАЯ РАБОТА"],
                    "task": ["ЗАДАНИЕ НА ВЫПУСКНУЮ", "ТЕХНИЧЕСКОЕ ЗАДАНИЕ"],
                    "calendar": ["КАЛЕНДАРНЫЙ ПЛАН"],
                    "definitions": ["ОПРЕДЕЛЕНИЯ", "ОБОЗНАЧЕНИЯ И СОКРАЩЕНИЯ", "ТЕРМИНЫ"]
                }
            }
        }


class VKRFormatter:
    """Комплексный форматтер ВКР согласно извлеченным требованиям"""

    ALIGN_MAP = {
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT
    }

    LINE_SPACING_MAP = {
        1.0: WD_LINE_SPACING.SINGLE,
        1.5: WD_LINE_SPACING.ONE_POINT_FIVE,
        2.0: WD_LINE_SPACING.DOUBLE
    }

    def __init__(self, requirements: Dict[str, Any]):
        self.requirements = requirements
        self.stats = {
            'total_paragraphs': 0,
            'h1_formatted': 0,
            'h2_formatted': 0,
            'lists_formatted': 0,
            'regular_formatted': 0,
            'skipped_sections': 0,
            'errors': 0
        }

    def format_document(self, input_path: str, output_path: str) -> bool:
        """Форматирует документ согласно требованиям"""

        try:
            doc = Document(input_path)

            # Применяем глобальные настройки
            self._apply_global_settings(doc)

            # Анализируем и форматируем параграфы
            self._format_all_paragraphs(doc)

            # Сохраняем документ
            doc.save(output_path)

            logger.info(f"Форматирование завершено. Статистика: {self.stats}")
            return True

        except Exception as e:
            logger.error(f"Ошибка форматирования: {e}")
            logger.error(traceback.format_exc())
            return False

    def _apply_global_settings(self, doc: Document) -> None:
        """Применяет глобальные настройки документа"""

        try:
            page_settings = self.requirements["global_formatting"]["page_settings"]
            margins = page_settings["margins_cm"]

            # Применяем поля ко всем разделам
            for section in doc.sections:
                section.top_margin = Cm(margins["top"])
                section.bottom_margin = Cm(margins["bottom"])
                section.left_margin = Cm(margins["left"])
                section.right_margin = Cm(margins["right"])

            logger.info("Глобальные настройки применены")

        except Exception as e:
            logger.error(f"Ошибка применения глобальных настроек: {e}")

    def _format_all_paragraphs(self, doc: Document) -> None:
        """Анализирует и форматирует все параграфы"""

        for paragraph in doc.paragraphs:
            self.stats['total_paragraphs'] += 1

            try:
                # Определяем тип параграфа
                paragraph_type = self._classify_paragraph(paragraph)

                # Применяем соответствующее форматирование
                if paragraph_type == "skip":
                    self.stats['skipped_sections'] += 1
                    continue
                elif paragraph_type == "h1":
                    self._format_h1(paragraph)
                    self.stats['h1_formatted'] += 1
                elif paragraph_type == "h2":
                    self._format_h2(paragraph)
                    self.stats['h2_formatted'] += 1
                elif paragraph_type == "list":
                    self._format_list_item(paragraph)
                    self.stats['lists_formatted'] += 1
                else:
                    self._format_regular_paragraph(paragraph)
                    self.stats['regular_formatted'] += 1

            except Exception as e:
                logger.warning(f"Ошибка форматирования параграфа: {e}")
                self.stats['errors'] += 1

    def _classify_paragraph(self, paragraph) -> str:
        """Классифицирует тип параграфа"""

        text = paragraph.text.strip()
        if not text:
            return "empty"

        # Проверяем, нужно ли пропустить секцию
        if self._should_skip_section(text):
            return "skip"

        # Проверяем заголовки H1
        if self._is_h1_header(text):
            return "h1"

        # Проверяем заголовки H2
        if self._is_h2_header(text):
            return "h2"

        # Проверяем элементы списка
        if self._is_list_item(text):
            return "list"

        return "regular"

    def _should_skip_section(self, text: str) -> bool:
        """Проверяет, нужно ли пропустить секцию"""

        skip_patterns = self.requirements["detection_rules"]["skip_sections"]

        for section_type, patterns in skip_patterns.items():
            for pattern in patterns:
                if pattern.upper() in text.upper():
                    logger.info(f"Пропускаем секцию: {text[:50]}...")
                    return True

        return False

    def _is_h1_header(self, text: str) -> bool:
        """Определяет заголовок H1"""

        patterns = self.requirements["detection_rules"]["h1_patterns"]

        for pattern in patterns:
            if re.match(pattern, text.upper().strip()):
                return True

        # Дополнительные проверки
        if len(text) < 100:
            upper_ratio = sum(1 for c in text if c.isupper()) / \
                len([c for c in text if c.isalpha()])
            if upper_ratio > 0.7:
                return True

        return False

    def _is_h2_header(self, text: str) -> bool:
        """Определяет заголовок H2"""

        patterns = self.requirements["detection_rules"]["h2_patterns"]

        for pattern in patterns:
            if re.match(pattern, text.strip()):
                return True

        return False

    def _is_list_item(self, text: str) -> bool:
        """Определяет элемент списка"""

        patterns = self.requirements["detection_rules"]["list_patterns"]

        for pattern in patterns:
            if re.match(pattern, text):
                return True

        return False

    def _format_h1(self, paragraph) -> None:
        """Форматирует заголовок H1"""

        h1_settings = self.requirements["headings"]["h1"]

        # Добавляем разрыв страницы если нужно
        if h1_settings.get("page_break", {}).get("before", False):
            if self._is_not_first_paragraph(paragraph):
                self._add_page_break(paragraph)

        # Форматируем текст
        self._apply_text_formatting(paragraph, h1_settings)

        # Применяем стили
        if h1_settings.get("font_style", {}).get("uppercase", False):
            self._make_uppercase(paragraph)

        # Выравнивание
        alignment = h1_settings.get("alignment", "center")
        paragraph.alignment = self.ALIGN_MAP.get(
            alignment, WD_ALIGN_PARAGRAPH.CENTER)

        # Отступы
        spacing = h1_settings.get("spacing", {})
        pf = paragraph.paragraph_format
        pf.space_before = Pt(spacing.get("before_pt", 18))
        pf.space_after = Pt(spacing.get("after_pt", 12))

        logger.debug(f"H1 отформатирован: {paragraph.text[:30]}...")

    def _format_h2(self, paragraph) -> None:
        """Форматирует заголовок H2"""

        h2_settings = self.requirements["headings"]["h2"]

        # Форматируем текст
        self._apply_text_formatting(paragraph, h2_settings)

        # Выравнивание
        alignment = h2_settings.get("alignment", "left")
        paragraph.alignment = self.ALIGN_MAP.get(
            alignment, WD_ALIGN_PARAGRAPH.LEFT)

        # Отступы
        spacing = h2_settings.get("spacing", {})
        pf = paragraph.paragraph_format
        pf.space_before = Pt(spacing.get("before_pt", 12))
        pf.space_after = Pt(spacing.get("after_pt", 6))

        logger.debug(f"H2 отформатирован: {paragraph.text[:30]}...")

    def _format_list_item(self, paragraph) -> None:
        """Форматирует элемент списка"""

        list_settings = self.requirements["lists"]["bullet_lists"]

        # Форматируем текст согласно настройкам списка
        font_settings = list_settings.get("font", {})
        self._apply_text_formatting(paragraph, {
            "font_name": font_settings.get("name", "Times New Roman"),
            "font_size": font_settings.get("size", 14)
        })

        # Выравнивание
        alignment = list_settings.get("alignment", "justify")
        paragraph.alignment = self.ALIGN_MAP.get(
            alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)

        # Отступ
        pf = paragraph.paragraph_format
        pf.left_indent = Cm(list_settings.get("indent_cm", 1.25))

        # Междустрочный интервал
        line_spacing = font_settings.get("line_spacing", 1.5)
        if line_spacing in self.LINE_SPACING_MAP:
            pf.line_spacing_rule = self.LINE_SPACING_MAP[line_spacing]

        logger.debug(f"Список отформатирован: {paragraph.text[:30]}...")

    def _format_regular_paragraph(self, paragraph) -> None:
        """Форматирует обычный параграф"""

        if not paragraph.text.strip():
            return

        base_text = self.requirements["global_formatting"]["base_text"]

        # Форматируем текст
        self._apply_text_formatting(paragraph, base_text)

        # Выравнивание
        alignment = base_text.get("alignment", "justify")
        paragraph.alignment = self.ALIGN_MAP.get(
            alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)

        # Отступ первой строки
        pf = paragraph.paragraph_format
        pf.first_line_indent = Cm(base_text.get("paragraph_indent_cm", 1.25))

        # Междустрочный интервал
        line_spacing = base_text.get("line_spacing", 1.5)
        if line_spacing in self.LINE_SPACING_MAP:
            pf.line_spacing_rule = self.LINE_SPACING_MAP[line_spacing]

    def _apply_text_formatting(self, paragraph, settings: Dict[str, Any]) -> None:
        """Применяет форматирование текста к параграфу"""

        if not paragraph.runs:
            paragraph.add_run()

        for run in paragraph.runs:
            font = run.font
            font.name = settings.get("font_name", "Times New Roman")
            font.size = Pt(settings.get("font_size", 14))

            # Жирность
            if settings.get("font_style", {}).get("bold", False):
                font.bold = True

    def _make_uppercase(self, paragraph) -> None:
        """Преобразует текст параграфа в верхний регистр"""

        original_text = paragraph.text
        paragraph.clear()
        paragraph.add_run(original_text.upper())

    def _add_page_break(self, paragraph) -> None:
        """Добавляет разрыв страницы перед параграфом"""

        if paragraph.runs:
            paragraph.runs[0].add_break(WD_BREAK.PAGE)
        else:
            paragraph.add_run().add_break(WD_BREAK.PAGE)

    def _is_not_first_paragraph(self, target_paragraph) -> bool:
        """Проверяет, не является ли параграф первым в документе"""

        try:
            # Получаем документ
            doc = target_paragraph._parent
            while hasattr(doc, '_parent') and doc._parent is not None:
                doc = doc._parent

            # Проверяем позицию
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph == target_paragraph:
                    # Проверяем, есть ли до него непустые параграфы
                    for j in range(i):
                        if doc.paragraphs[j].text.strip():
                            return True
                    return False

            return True

        except Exception:
            return True

    def get_stats(self) -> Dict[str, int]:
        """Возвращает статистику форматирования"""
        return self.stats.copy()


# Главная функция для API
def comprehensive_format_vkr(vkr_path: str, requirements_path: str, output_path: str) -> Tuple[bool, Dict[str, Any]]:
    """
    Комплексное форматирование ВКР

    Returns:
        Tuple[bool, Dict]: (успех, статистика)
    """
    try:
        # Анализируем требования
        analyzer = VKRAnalyzer()
        requirements = analyzer.analyze_requirements(requirements_path)

        # Форматируем документ
        formatter = VKRFormatter(requirements)
        success = formatter.format_document(vkr_path, output_path)

        stats = formatter.get_stats()
        stats['requirements_extracted'] = True

        return success, stats

    except Exception as e:
        logger.error(f"Ошибка комплексного форматирования: {e}")
        return False, {'error': str(e)}
