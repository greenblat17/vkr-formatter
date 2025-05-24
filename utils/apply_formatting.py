from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
import logging
from typing import Dict, Any, Optional, List
import traceback
import re

# Настройка логирования
logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class VKRFormatter:
    """Класс для форматирования ВКР согласно требованиям (для заголовков 1 и 2 уровня)"""
    
    ALIGN_MAP = {
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "по ширине": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "по левому краю": WD_ALIGN_PARAGRAPH.LEFT,
        "по центру": WD_ALIGN_PARAGRAPH.CENTER,
        "по правому краю": WD_ALIGN_PARAGRAPH.RIGHT
    }
    
    def __init__(self):
        self.stats = {
            'total_paragraphs': 0,
            'h1_found': 0,
            'h1_processed': 0,
            'h2_found': 0,
            'h2_processed': 0,
            'title_page_paragraphs': 0,
            'errors': 0
        }
        # Счетчики для нумерации
        self.current_h1_number = 0
        self.current_h2_number = 0
    
    def _is_title_page(self, paragraph) -> bool:
        """
        Определяет, является ли параграф частью титульного листа
        """
        # Проверяем стиль
        if paragraph.style and paragraph.style.name:
            style_name = paragraph.style.name.lower()
            if any(keyword in style_name for keyword in ['title', 'титул', 'cover']):
                return True
        
        # Проверяем содержимое
        text = paragraph.text.strip().lower()
        title_keywords = [
            'министерство', 'образования', 'науки', 'федеральное', 'государственное',
            'бюджетное', 'образовательное', 'учреждение', 'высшего', 'образования',
            'дипломная', 'работа', 'выполнил', 'студент', 'группы', 'научный',
            'руководитель', 'должность', 'ученая', 'степень', 'город', 'год'
        ]
        
        # Если параграф содержит ключевые слова титульного листа
        if any(keyword in text for keyword in title_keywords):
            return True
            
        # Проверяем форматирование
        if paragraph.runs:
            for run in paragraph.runs:
                # Если текст отцентрирован и имеет большой размер шрифта
                if (paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER and 
                    run.font.size and run.font.size.pt > 14):
                    return True
        
        return False
    
    def apply_formatting(self, input_path: str, formatting: Dict[str, Any], output_path: str) -> bool:
        """
        Применяет форматирование к заголовкам 1 и 2 уровня, пропуская титульный лист
        
        Args:
            input_path: путь к исходному документу
            formatting: словарь с требованиями к форматированию
            output_path: путь к выходному документу
            
        Returns:
            bool: успешность операции
        """
        try:
            # Сбрасываем статистику и счетчики
            self.stats = {
                'total_paragraphs': 0,
                'h1_found': 0,
                'h1_processed': 0,
                'h2_found': 0,
                'h2_processed': 0,
                'title_page_paragraphs': 0,
                'errors': 0
            }
            self.current_h1_number = 0
            self.current_h2_number = 0
            
            doc = Document(input_path)
            
            # Сначала проходим по документу и определяем структуру
            self._analyze_document_structure(doc)
            
            # Затем форматируем заголовки
            self._format_headers(doc, formatting)
            
            # Сохраняем документ
            doc.save(output_path)
            
            logger.info(f"Форматирование заголовков завершено. Статистика: {self.stats}")
            return True
            
        except Exception as e:
            logger.error(f"Ошибка при форматировании: {str(e)}")
            logger.error(traceback.format_exc())
            return False
    
    def _analyze_document_structure(self, doc: Document) -> None:
        """Анализирует структуру документа для определения правильной нумерации"""
        # Сбрасываем счетчики
        self.current_h1_number = 0
        self.current_h2_number = 0
        
        # Словарь для хранения обнаруженных заголовков с их порядковыми номерами
        self.headers_map = {}
        
        # Проходим по всем параграфам и определяем структуру
        h1_counter = 0
        h2_counter = 0
        last_h1_index = -1
        
        for i, paragraph in enumerate(doc.paragraphs):
            if self._is_title_page(paragraph):
                continue
            
            if self._is_h1_header(paragraph):
                h1_counter += 1
                last_h1_index = i
                h2_counter = 0  # Сбрасываем счетчик h2 при обнаружении нового h1
                self.headers_map[i] = {'type': 'h1', 'number': h1_counter, 'parent': None}
                logger.debug(f"Анализ: H1 #{h1_counter} найден в позиции {i}")
            
            elif self._is_h2_header(paragraph):
                h2_counter += 1
                self.headers_map[i] = {'type': 'h2', 'number': h2_counter, 'parent': last_h1_index}
                logger.debug(f"Анализ: H2 #{h1_counter}.{h2_counter} найден в позиции {i}")
        
        logger.info(f"Анализ структуры документа: обнаружено {h1_counter} заголовков H1 и несколько H2")
    
    def _format_headers(self, doc: Document, formatting: Dict[str, Any]) -> None:
        """Ищет и форматирует заголовки 1 и 2 уровня, пропуская титульный лист"""
        
        for i, paragraph in enumerate(doc.paragraphs):
            self.stats['total_paragraphs'] += 1
            
            try:
                # Пропускаем параграфы титульного листа
                if self._is_title_page(paragraph):
                    self.stats['title_page_paragraphs'] += 1
                    logger.debug(f"Пропущен параграф титульного листа: '{paragraph.text[:50]}...'")
                    continue
                
                # Проверяем на H1
                if self._is_h1_header(paragraph):
                    self.stats['h1_found'] += 1
                    
                    # Получаем номер заголовка из структуры
                    header_info = self.headers_map.get(i, {})
                    h1_number = header_info.get('number', 0)
                    
                    logger.info(f"Найден H1 #{h1_number}: '{paragraph.text[:50]}...'")
                    
                    # Применяем форматирование к H1
                    self._format_h1_paragraph(paragraph, formatting, h1_number)
                    self.stats['h1_processed'] += 1
                    continue
                
                # Проверяем на H2
                if self._is_h2_header(paragraph):
                    self.stats['h2_found'] += 1
                    
                    # Получаем номер заголовка и родителя из структуры
                    header_info = self.headers_map.get(i, {})
                    h2_number = header_info.get('number', 0)
                    parent_index = header_info.get('parent', -1)
                    
                    # Получаем номер родительского H1
                    parent_info = self.headers_map.get(parent_index, {})
                    h1_number = parent_info.get('number', 0)
                    
                    logger.info(f"Найден H2 #{h1_number}.{h2_number}: '{paragraph.text[:50]}...'")
                    
                    # Применяем форматирование к H2
                    self._format_h2_paragraph(paragraph, formatting, h1_number, h2_number)
                    self.stats['h2_processed'] += 1
                    
            except Exception as e:
                logger.warning(f"Ошибка при обработке параграфа: {e}")
                self.stats['errors'] += 1
                continue
    
    def _is_h1_header(self, paragraph) -> bool:
        """
        Определяет, является ли параграф заголовком 1 уровня
        Используем несколько критериев:
        1. Стиль параграфа содержит "Heading 1" или "Заголовок 1"
        2. Текст короткий (меньше 100 символов) и написан заглавными буквами
        3. Параграф имеет больший размер шрифта
        """
        
        # Критерий 1: Проверяем стиль
        if paragraph.style and paragraph.style.name:
            style_name = paragraph.style.name.lower()
            if any(keyword in style_name for keyword in ['heading 1', 'заголовок 1', 'title']):
                logger.debug(f"H1 найден по стилю: {paragraph.style.name}")
                return True
        
        return False
    
    def _is_h2_header(self, paragraph) -> bool:
        """
        Определяет, является ли параграф заголовком 2 уровня
        Используем несколько критериев:
        1. Стиль параграфа содержит "Heading 2" или "Заголовок 2"
        2. Текст начинается с цифр в формате "1.1." или "1.1"
        3. Параграф имеет средний размер шрифта
        """
        # Критерий 1: Проверяем стиль
        if paragraph.style and paragraph.style.name:
            style_name = paragraph.style.name.lower()
            if any(keyword in style_name for keyword in ['heading 2', 'заголовок 2', 'subtitle']):
                logger.debug(f"H2 найден по стилю: {paragraph.style.name}")
                return True
        
        return False
    
    def _format_h1_paragraph(self, paragraph, formatting: Dict[str, Any], h1_number: int = 0) -> None:
        """Применяет форматирование к заголовку 1 уровня"""
        
        # Получаем настройки для H1 из форматирования
        h1_settings = formatting.get("h1_formatting", {})
        
        # Если специальных настроек для H1 нет, используем общие
        if not h1_settings:
            h1_settings = {
                "font_name": formatting.get("font_name", "Times New Roman"),
                "font_size": formatting.get("font_size_h1", formatting.get("font_size_main", 16)),
                "alignment": formatting.get("h1_alignment", "center"),
                "bold": formatting.get("h1_bold", True),
                "uppercase": formatting.get("h1_uppercase", True),
                "space_before": formatting.get("h1_space_before", 12),
                "space_after": formatting.get("h1_space_after", 12),
                "numbering": formatting.get("h1_numbering", "1.")
            }
        
        # Применяем нумерацию, заменяя шаблон
        numbering_template = h1_settings.get("numbering", "")
        if numbering_template and h1_number > 0:
            # Заменяем цифру в шаблоне на актуальный номер
            actual_numbering = numbering_template.replace("1", str(h1_number))
            paragraph.text = self._apply_numbering(paragraph.text, actual_numbering)
        
        # Форматируем runs
        if not paragraph.runs:
            paragraph.add_run()
        
        # Приводим к верхнему регистру если нужно
        if h1_settings.get("uppercase", True):
            paragraph.text = paragraph.text.upper()
        
        for run in paragraph.runs:
            font = run.font
            
            # Шрифт и размер
            font.name = h1_settings.get("font_name", "Times New Roman")
            font.size = Pt(h1_settings.get("font_size", 16))
            
            # Жирность
            if h1_settings.get("bold", True):
                font.bold = True
        
        # Выравнивание
        alignment = h1_settings.get("alignment", "center")
        paragraph.alignment = self.ALIGN_MAP.get(alignment, WD_ALIGN_PARAGRAPH.CENTER)
        
        # Отступы и интервалы
        pf = paragraph.paragraph_format
        pf.space_before = Pt(h1_settings.get("space_before", 12))
        pf.space_after = Pt(h1_settings.get("space_after", 12))
        
        logger.info(f"H1 #{h1_number} отформатирован: {paragraph.text[:30]}...")
    
    def _format_h2_paragraph(self, paragraph, formatting: Dict[str, Any], h1_number: int = 0, h2_number: int = 0) -> None:
        """Применяет форматирование к заголовку 2 уровня"""
        
        # Получаем настройки для H2 из форматирования
        h2_settings = formatting.get("h2_formatting", {})
        
        # Если специальных настроек для H2 нет, используем общие
        if not h2_settings:
            h2_settings = {
                "font_name": formatting.get("font_name", "Times New Roman"),
                "font_size": formatting.get("font_size_h2", formatting.get("font_size_main", 14)),
                "alignment": formatting.get("h2_alignment", "left"),
                "bold": formatting.get("h2_bold", True),
                "uppercase": formatting.get("h2_uppercase", False),
                "space_before": formatting.get("h2_space_before", 12),
                "space_after": formatting.get("h2_space_after", 12),
                "numbering": formatting.get("h2_numbering", "1.1.")
            }
        
        # Применяем нумерацию, заменяя шаблон
        numbering_template = h2_settings.get("numbering", "")
        if numbering_template and h1_number > 0 and h2_number > 0:
            # Заменяем цифры в шаблоне на актуальные номера
            actual_numbering = numbering_template.replace("1.1", f"{h1_number}.{h2_number}")
            if actual_numbering == numbering_template:  # Если замена не произошла
                actual_numbering = numbering_template.replace("1", str(h1_number))
                actual_numbering = actual_numbering.replace("1", str(h2_number), 1)
            paragraph.text = self._apply_numbering(paragraph.text, actual_numbering)
        
        # Форматируем runs
        if not paragraph.runs:
            paragraph.add_run()
        
        # Приводим к верхнему регистру если нужно
        if h2_settings.get("uppercase", False):
            paragraph.text = paragraph.text.upper()
        
        for run in paragraph.runs:
            font = run.font
            
            # Шрифт и размер
            font.name = h2_settings.get("font_name", "Times New Roman")
            font.size = Pt(h2_settings.get("font_size", 14))
            
            # Жирность
            if h2_settings.get("bold", True):
                font.bold = True
        
        # Выравнивание
        alignment = h2_settings.get("alignment", "left")
        paragraph.alignment = self.ALIGN_MAP.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)
        
        # Отступы и интервалы
        pf = paragraph.paragraph_format
        pf.space_before = Pt(h2_settings.get("space_before", 12))
        pf.space_after = Pt(h2_settings.get("space_after", 12))
        
        logger.info(f"H2 #{h1_number}.{h2_number} отформатирован: {paragraph.text[:30]}...")
    
    def validate_formatting(self, formatting: Dict[str, Any]) -> Dict[str, Any]:
        """Валидирует форматирование для H1 и H2"""
        validated = {}
        
        # Основные настройки
        validated["font_name"] = formatting.get("font_name", "Times New Roman")
        validated["font_size_main"] = max(8, min(72, formatting.get("font_size_main", 14)))
        
        # Настройки для H1
        h1_formatting = formatting.get("h1_formatting", {})
        validated["h1_formatting"] = {
            "font_name": h1_formatting.get("font_name", validated["font_name"]),
            "font_size": max(8, min(72, h1_formatting.get("font_size", 16))),
            "alignment": h1_formatting.get("alignment", "center"),
            "bold": h1_formatting.get("bold", True),
            "uppercase": h1_formatting.get("uppercase", True),
            "space_before": max(0, h1_formatting.get("space_before", 12)),
            "space_after": max(0, h1_formatting.get("space_after", 12)),
            "numbering": h1_formatting.get("numbering", "1.")
        }
        
        # Настройки для H2
        h2_formatting = formatting.get("h2_formatting", {})
        validated["h2_formatting"] = {
            "font_name": h2_formatting.get("font_name", validated["font_name"]),
            "font_size": max(8, min(72, h2_formatting.get("font_size", 14))),
            "alignment": h2_formatting.get("alignment", "left"),
            "bold": h2_formatting.get("bold", True),
            "uppercase": h2_formatting.get("uppercase", False),
            "space_before": max(0, h2_formatting.get("space_before", 12)),
            "space_after": max(0, h2_formatting.get("space_after", 12)),
            "numbering": h2_formatting.get("numbering", "1.1.")
        }
        
        return validated
    
    def get_stats(self) -> Dict[str, int]:
        """Возвращает статистику обработки"""
        return self.stats.copy()

    def _apply_numbering(self, text, numbering):
        """Применяет нумерацию к тексту, учитывая текущий формат"""
        if not numbering:
            return text
        
        # Удаляем пробелы в начале
        text = text.lstrip()
        
        # Если уже начинается с нужной нумерации — ничего не делаем
        if text.startswith(numbering):
            return text
        
        # Если начинается с другой нумерации — заменяем
        pattern = r'^\d+(\.\d+)*\.?\s*'
        if re.match(pattern, text):
            text_wo_numbering = re.sub(pattern, '', text)
            return f"{numbering} {text_wo_numbering.lstrip()}"
        
        # Если нумерации нет — просто добавляем
        return f"{numbering} {text}"


def apply_formatting(input_path: str, formatting: Dict[str, Any], output_path: str) -> bool:
    """
    Функция-обертка для обратной совместимости
    """
    formatter = VKRFormatter()
    validated_formatting = formatter.validate_formatting(formatting)
    return formatter.apply_formatting(input_path, validated_formatting, output_path)


    # formatter = VKRFormatter()
    # success = formatter.apply_formatting("input.docx", sample_formatting, "output.docx")
    # print(f"Форматирование H1 {'успешно' if success else 'не удалось'}")
    # print(f"Статистика: {formatter.get_stats()}")