from docx import Document
from openai import OpenAI
import json
from dotenv import load_dotenv
import os
import logging
from typing import Dict, Any, List
import traceback

# Load environment variables
load_dotenv()

# Get API key from environment variable
api_key = os.getenv('OPENAI_API_KEY')
if not api_key:
    raise ValueError("OPENAI_API_KEY not found in environment variables")

client = OpenAI(api_key=api_key)

# Настройка логирования
logging.basicConfig(level=logging.INFO,
                    format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)


class AIDirectFormatter:
    """Класс для прямого форматирования документов через ИИ"""

    def __init__(self):
        self.stats = {
            'paragraphs_analyzed': 0,
            'h1_formatted': 0,
            'regular_formatted': 0,
            'errors': 0
        }

    def format_document_with_ai(self, vkr_path: str, requirements_path: str, output_path: str) -> bool:
        """
        Форматирует документ, используя ИИ для анализа требований и текста

        Args:
            vkr_path: путь к файлу ВКР
            requirements_path: путь к файлу с требованиями
            output_path: путь к выходному файлу

        Returns:
            bool: успешность операции
        """
        try:
            # Читаем оба документа
            vkr_doc = Document(vkr_path)
            req_doc = Document(requirements_path)

            # Извлекаем текст требований
            requirements_text = "\n".join(
                [p.text for p in req_doc.paragraphs if p.text.strip()])

            # Получаем инструкции по форматированию от ИИ
            formatting_instructions = self._get_ai_formatting_instructions(
                requirements_text)

            # Применяем форматирование к каждому параграфу через ИИ
            self._apply_ai_formatting(vkr_doc, formatting_instructions)

            # Сохраняем документ
            vkr_doc.save(output_path)

            logger.info(
                f"ИИ форматирование завершено. Статистика: {self.stats}")
            return True

        except Exception as e:
            logger.error(f"Ошибка при ИИ форматировании: {str(e)}")
            logger.error(traceback.format_exc())
            return False

    def _get_ai_formatting_instructions(self, requirements_text: str) -> Dict[str, Any]:
        """Получает инструкции по форматированию от ИИ"""

        prompt = f"""Проанализируй требования к оформлению дипломной работы и дай конкретные инструкции по форматированию:

ТРЕБОВАНИЯ:
---
{requirements_text}
---

Выдай JSON с инструкциями для:
1. Глобальных настроек документа (поля, основной шрифт, отступы)
2. Заголовков 1 уровня (как их форматировать)
3. Правил определения заголовков 1 уровня (как их найти в тексте)

Будь конкретным и точным."""

        function_schema = {
            "name": "formatting_instructions",
            "description": "Получение детальных инструкций по форматированию от ИИ",
            "parameters": {
                "type": "object",
                "properties": {
                    "global_formatting": {
                        "type": "object",
                        "properties": {
                            "margins_cm": {
                                "type": "object",
                                "properties": {
                                    "top": {"type": "number"},
                                    "bottom": {"type": "number"},
                                    "left": {"type": "number"},
                                    "right": {"type": "number"}
                                }
                            },
                            "main_font": {"type": "string"},
                            "main_font_size": {"type": "integer"},
                            "line_spacing": {"type": "number"},
                            "paragraph_indent_cm": {"type": "number"},
                            "text_alignment": {"type": "string"}
                        }
                    },
                    "h1_formatting": {
                        "type": "object",
                        "properties": {
                            "font_name": {"type": "string"},
                            "font_size": {"type": "integer"},
                            "alignment": {"type": "string"},
                            "bold": {"type": "boolean"},
                            "uppercase": {"type": "boolean"},
                            "new_page_before": {"type": "boolean"},
                            "space_before_pt": {"type": "integer"},
                            "space_after_pt": {"type": "integer"}
                        }
                    },
                    "h1_detection_rules": {
                        "type": "object",
                        "properties": {
                            "patterns": {
                                "type": "array",
                                "items": {"type": "string"},
                                "description": "Регулярные выражения для поиска H1"
                            },
                            "style_keywords": {
                                "type": "array",
                                "items": {"type": "string"},
                                "description": "Ключевые слова в стилях"
                            },
                            "max_length": {"type": "integer"},
                            "uppercase_ratio": {"type": "number"}
                        }
                    }
                },
                "required": ["global_formatting", "h1_formatting", "h1_detection_rules"]
            }
        }

        try:
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                tools=[{"type": "function", "function": function_schema}],
                tool_choice="required",
            )

            args = response.choices[0].message.tool_calls[0].function.arguments
            instructions = json.loads(args)

            logger.info("ИИ инструкции получены")
            return instructions

        except Exception as e:
            logger.error(f"Ошибка получения ИИ инструкций: {e}")
            # Возвращаем дефолтные инструкции
            return self._get_default_instructions()

    def _apply_ai_formatting(self, doc: Document, instructions: Dict[str, Any]) -> None:
        """Применяет форматирование согласно ИИ инструкциям"""

        # Применяем глобальные настройки
        self._apply_global_formatting(doc, instructions["global_formatting"])

        # Форматируем параграфы
        for paragraph in doc.paragraphs:
            self.stats['paragraphs_analyzed'] += 1

            try:
                # Спрашиваем у ИИ, как форматировать этот конкретный параграф
                paragraph_format = self._get_paragraph_formatting_from_ai(
                    paragraph.text,
                    instructions
                )

                if paragraph_format["is_h1"]:
                    self._format_as_h1(
                        paragraph, instructions["h1_formatting"])
                    self.stats['h1_formatted'] += 1
                else:
                    self._format_as_regular(
                        paragraph, instructions["global_formatting"])
                    self.stats['regular_formatted'] += 1

            except Exception as e:
                logger.warning(f"Ошибка форматирования параграфа: {e}")
                self.stats['errors'] += 1

    def _get_paragraph_formatting_from_ai(self, paragraph_text: str, instructions: Dict[str, Any]) -> Dict[str, Any]:
        """Спрашивает у ИИ, как форматировать конкретный параграф"""

        # Для коротких запросов можем использовать более простую логику
        if not paragraph_text.strip() or len(paragraph_text) > 200:
            return {"is_h1": False, "formatting": "regular"}

        prompt = f"""Определи, является ли этот текст заголовком 1 уровня согласно правилам:

ТЕКСТ: "{paragraph_text}"

ПРАВИЛА ОПРЕДЕЛЕНИЯ H1:
{json.dumps(instructions["h1_detection_rules"], ensure_ascii=False, indent=2)}

Ответь только JSON: {{"is_h1": true/false, "confidence": 0.0-1.0, "reason": "почему"}}"""

        try:
            response = client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=100,
                temperature=0.1
            )

            result = json.loads(response.choices[0].message.content)
            return result

        except Exception as e:
            logger.warning(f"Ошибка анализа параграфа ИИ: {e}")
            # Fallback на простую логику
            return {"is_h1": self._simple_h1_detection(paragraph_text)}

    def _simple_h1_detection(self, text: str) -> bool:
        """Простая логика определения H1 как fallback"""
        if not text or len(text) > 100:
            return False

        import re
        # Проверяем паттерны
        patterns = [
            r'^\d+\.\s*[А-ЯЁ\s]+$',  # "1. ВВЕДЕНИЕ"
            r'^ГЛАВА\s+\d+',          # "ГЛАВА 1"
            r'^\d+\.$',               # "1."
            r'^[IVX]+\.\s*[А-ЯЁ\s]+$'  # "I. ВВЕДЕНИЕ"
        ]

        for pattern in patterns:
            if re.match(pattern, text.upper().strip()):
                return True

        # Проверяем процент заглавных букв
        if text:
            upper_ratio = sum(1 for c in text if c.isupper()) / \
                len([c for c in text if c.isalpha()])
            return upper_ratio > 0.7 and len(text) < 50

        return False

    def _apply_global_formatting(self, doc: Document, global_format: Dict[str, Any]) -> None:
        """Применяет глобальные настройки"""
        try:
            # Настройки полей
            margins = global_format.get("margins_cm", {})
            for section in doc.sections:
                from docx.shared import Cm
                section.top_margin = Cm(margins.get("top", 2.0))
                section.bottom_margin = Cm(margins.get("bottom", 2.0))
                section.left_margin = Cm(margins.get("left", 3.0))
                section.right_margin = Cm(margins.get("right", 1.5))

            logger.info("Глобальные настройки применены")

        except Exception as e:
            logger.error(f"Ошибка применения глобальных настроек: {e}")

    def _format_as_h1(self, paragraph, h1_format: Dict[str, Any]) -> None:
        """Форматирует параграф как заголовок H1"""
        try:
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK

            # Добавляем разрыв страницы если нужно
            if h1_format.get("new_page_before", False):
                if paragraph.runs:
                    paragraph.runs[0].add_break(WD_BREAK.PAGE)
                else:
                    paragraph.add_run().add_break(WD_BREAK.PAGE)

            # Форматируем текст
            if not paragraph.runs:
                paragraph.add_run()

            for run in paragraph.runs:
                font = run.font
                font.name = h1_format.get("font_name", "Times New Roman")
                font.size = Pt(h1_format.get("font_size", 16))
                font.bold = h1_format.get("bold", True)

            # Заглавные буквы
            if h1_format.get("uppercase", True):
                original_text = paragraph.text
                paragraph.clear()
                paragraph.add_run(original_text.upper())

            # Выравнивание
            alignment_map = {
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
                "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
            }
            paragraph.alignment = alignment_map.get(h1_format.get(
                "alignment", "center"), WD_ALIGN_PARAGRAPH.CENTER)

            # Отступы
            pf = paragraph.paragraph_format
            pf.space_before = Pt(h1_format.get("space_before_pt", 12))
            pf.space_after = Pt(h1_format.get("space_after_pt", 12))

        except Exception as e:
            logger.error(f"Ошибка форматирования H1: {e}")

    def _format_as_regular(self, paragraph, global_format: Dict[str, Any]) -> None:
        """Форматирует параграф как обычный текст"""
        try:
            from docx.shared import Pt, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

            if not paragraph.text.strip():
                return

            if not paragraph.runs:
                paragraph.add_run()

            # Форматируем runs
            for run in paragraph.runs:
                font = run.font
                font.name = global_format.get("main_font", "Times New Roman")
                font.size = Pt(global_format.get("main_font_size", 14))

            # Форматирование параграфа
            pf = paragraph.paragraph_format

            # Выравнивание
            alignment_map = {
                "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
                "center": WD_ALIGN_PARAGRAPH.CENTER
            }
            paragraph.alignment = alignment_map.get(global_format.get(
                "text_alignment", "justify"), WD_ALIGN_PARAGRAPH.JUSTIFY)

            # Отступ первой строки
            indent = global_format.get("paragraph_indent_cm", 1.25)
            pf.first_line_indent = Cm(indent)

            # Междустрочный интервал
            line_spacing = global_format.get("line_spacing", 1.5)
            if line_spacing == 1.0:
                pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
            elif line_spacing == 1.5:
                pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            elif line_spacing == 2.0:
                pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            else:
                pf.line_spacing = line_spacing

        except Exception as e:
            logger.error(f"Ошибка форматирования обычного текста: {e}")

    def _get_default_instructions(self) -> Dict[str, Any]:
        """Возвращает инструкции по умолчанию"""
        return {
            "global_formatting": {
                "margins_cm": {"top": 2.0, "bottom": 2.0, "left": 3.0, "right": 1.5},
                "main_font": "Times New Roman",
                "main_font_size": 14,
                "line_spacing": 1.5,
                "paragraph_indent_cm": 1.25,
                "text_alignment": "justify"
            },
            "h1_formatting": {
                "font_name": "Times New Roman",
                "font_size": 16,
                "alignment": "center",
                "bold": True,
                "uppercase": True,
                "new_page_before": True,
                "space_before_pt": 12,
                "space_after_pt": 12
            },
            "h1_detection_rules": {
                "patterns": [r'^\d+\.\s*[А-ЯЁ\s]+$', r'^ГЛАВА\s+\d+'],
                "style_keywords": ["heading 1", "заголовок 1"],
                "max_length": 100,
                "uppercase_ratio": 0.7
            }
        }

    def get_stats(self) -> Dict[str, int]:
        """Возвращает статистику"""
        return self.stats.copy()


# Функция для интеграции с существующим API
def format_with_ai(vkr_path: str, requirements_path: str, output_path: str) -> bool:
    """
    Функция-обертка для использования в существующем API
    """
    formatter = AIDirectFormatter()
    return formatter.format_document_with_ai(vkr_path, requirements_path, output_path)
