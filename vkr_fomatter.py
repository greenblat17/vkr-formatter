from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
import re
import logging
from typing import Dict, Any, List
from pathlib import Path
import colorlog

# Настройка цветного логирования
def setup_colored_logging():
    """Настраивает цветное логирование"""
    
    # Создаем цветной форматтер
    color_formatter = colorlog.ColoredFormatter(
        '%(log_color)s%(asctime)s - %(levelname)-8s%(reset)s %(message)s',
        datefmt='%H:%M:%S',
        log_colors={
            'DEBUG': 'cyan',
            'INFO': 'green', 
            'WARNING': 'yellow',
            'ERROR': 'red',
            'CRITICAL': 'red,bg_white',
        },
        secondary_log_colors={},
        style='%'
    )
    
    # Настраиваем handler
    handler = colorlog.StreamHandler()
    handler.setFormatter(color_formatter)
    
    # Настраиваем логгер
    logger = colorlog.getLogger(__name__)
    logger.setLevel(logging.INFO)
    logger.handlers.clear()
    logger.addHandler(handler)
    
    return logger

logger = setup_colored_logging()

class SimpleVKRFormatter:
    """Простой и понятный форматтер ВКР"""
    
    def __init__(self, requirements: Dict[str, Any]):
        """
        Args:
            requirements: словарь требований из заглушки
        """
        self.requirements = requirements
        self.stats = {
            'total_paragraphs': 0,
            'skipped_paragraphs': 0,
            'h1_formatted': 0,
            'h2_formatted': 0,
            'lists_formatted': 0,
            'regular_formatted': 0,
            'errors': 0
        }
        
        # Состояние обработки документа
        self.document_state = {
            'in_title_section': True,  # Начинаем с предположения, что мы в титульной секции
            'in_contents_section': False,  # Находимся ли в разделе содержания
            'found_main_content': False,  # Нашли ли основное содержание
            'pages_skipped': 0  # Количество пропущенных страниц
        }
        
        # Маппинги для удобства
        self.align_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        
        self.line_spacing_map = {
            1.0: WD_LINE_SPACING.SINGLE,
            1.5: WD_LINE_SPACING.ONE_POINT_FIVE,
            2.0: WD_LINE_SPACING.DOUBLE
        }
    
    def format_document(self, input_path: str, output_path: str) -> bool:
        """
        Основная функция форматирования документа
        
        Args:
            input_path: путь к исходному документу
            output_path: путь к результату
            
        Returns:
            bool: успешность операции
        """
        try:
            logger.info(f"📂 Начинаем форматирование: {input_path}")
            logger.info(f"💾 Выходной путь: {output_path}")
            
            # Проверяем входной файл
            input_file = Path(input_path)
            if not input_file.exists():
                logger.error(f"❌ Входной файл не существует: {input_path}")
                return False
            
            # Загружаем документ
            logger.info("📖 Загружаем документ...")
            doc = Document(input_path)
            logger.info(f"✅ Документ загружен, параграфов: {len(doc.paragraphs)}")
            
            # Шаг 1: Применяем глобальные настройки (поля, базовый шрифт)
            logger.info("⚙️  Применяем глобальные настройки...")
            self._apply_global_settings(doc)
            
            # Шаг 2: Обрабатываем каждый параграф
            logger.info("🔄 Обрабатываем параграфы...")
            self._process_all_paragraphs(doc)
            
            # Шаг 3: Сохраняем результат
            logger.info(f"💾 Сохраняем документ в: {output_path}")
            doc.save(output_path)
            
            # Проверяем, что файл создался
            output_file = Path(output_path)
            if output_file.exists():
                logger.info(f"✅ Файл создан, размер: {output_file.stat().st_size} байт")
            else:
                logger.error(f"❌ Файл НЕ создался: {output_path}")
                return False
            
            logger.info(f"🎉 Форматирование завершено! Статистика: {self.stats}")
            return True
            
        except Exception as e:
            logger.error(f"Ошибка форматирования: {e}")
            import traceback
            logger.error(f"Полная трассировка: {traceback.format_exc()}")
            return False
    
    def _apply_global_settings(self, doc: Document) -> None:
        """Применяет глобальные настройки: поля страницы"""
        
        try:
            margins = self.requirements["base_formatting"]["margins_cm"]
            
            # Применяем поля ко всем секциям
            for section in doc.sections:
                section.top_margin = Cm(margins["top"])
                section.bottom_margin = Cm(margins["bottom"])
                section.left_margin = Cm(margins["left"])
                section.right_margin = Cm(margins["right"])
            
            logger.info(f"Применены поля: {margins}")
            
        except Exception as e:
            logger.error(f"Ошибка применения глобальных настроек: {e}")
            self.stats['errors'] += 1
    
    def _process_all_paragraphs(self, doc: Document) -> None:
        """Обрабатывает все параграфы документа"""
        
        logger.info("Начинаем обработку параграфов...")
        
        for i, paragraph in enumerate(doc.paragraphs):
            self.stats['total_paragraphs'] += 1
            
            try:
                text = paragraph.text.strip()
                
                # Определяем тип параграфа
                paragraph_type = self._classify_paragraph(text)
                
                # Логируем каждый параграф для отладки
                if text:  # Логируем только непустые
                    logger.debug(f"Параграф {i+1}: тип='{paragraph_type}', текст='{text[:100]}{'...' if len(text) > 100 else ''}'")
                
                # Применяем соответствующее форматирование
                if paragraph_type == "skip":
                    self.stats['skipped_paragraphs'] += 1
                    logger.info(f"⏭️  ПРОПУСК #{i+1}: {text[:60]}{'...' if len(text) > 60 else ''}")
                    
                elif paragraph_type == "h1":
                    self._format_h1_paragraph(paragraph)
                    self.stats['h1_formatted'] += 1
                    logger.info(f"📝 H1 #{i+1}: {text[:40]}...")
                    
                elif paragraph_type == "h2":
                    self._format_h2_paragraph(paragraph)
                    self.stats['h2_formatted'] += 1
                    logger.info(f"📄 H2 #{i+1}: {text[:40]}...")
                    
                elif paragraph_type == "list":
                    self._format_list_paragraph(paragraph)
                    self.stats['lists_formatted'] += 1
                    logger.debug(f"📋 СПИСОК #{i+1}: {text[:40]}...")
                    
                else:  # regular
                    self._format_regular_paragraph(paragraph)
                    self.stats['regular_formatted'] += 1
                    
            except Exception as e:
                logger.warning(f"Ошибка обработки параграфа {i+1}: {e}")
                self.stats['errors'] += 1
        
        logger.info(f"Обработка параграфов завершена. Статистика: {self.stats}")
    
    def _classify_paragraph(self, text: str) -> str:
        """
        Классифицирует тип параграфа с учетом состояния документа
        
        Returns:
            str: "skip", "h1", "h2", "list", "regular"
        """
        if not text:
            return "skip"
        
        text_clean = text.strip()
        
        # 1. Проверяем заголовок содержания
        if self._is_contents_header(text_clean):
            logger.info(f"📑 НАЧАЛО СОДЕРЖАНИЯ: {text_clean[:50]}...")
            self.document_state['in_contents_section'] = True
            return "skip"
        
        # 2. Если мы в разделе содержания, проверяем не началось ли основное содержание
        if self.document_state['in_contents_section']:
            # Если это строка содержания - продолжаем пропускать
            if self._is_contents_line(text_clean):
                logger.debug(f"📑 СОДЕРЖАНИЕ (строка): {text_clean[:50]}...")
                return "skip"
            
            # Если это НЕ строка содержания и похоже на начало основного содержания
            if self._is_main_content_start(text_clean):
                logger.info(f"🟢 КОНЕЦ СОДЕРЖАНИЯ, НАЧАЛО ОСНОВНОГО СОДЕРЖАНИЯ: {text_clean[:60]}...")
                self.document_state['in_contents_section'] = False
                self.document_state['in_title_section'] = False
                self.document_state['found_main_content'] = True
                
                # Определяем тип этого параграфа
                if self._is_h1_paragraph_content(text_clean):
                    return "h1"
                elif self._is_h2_paragraph_content(text_clean):
                    return "h2"
                else:
                    return "regular"
            else:
                # Пустая строка или что-то неопределенное в содержании - пропускаем
                logger.debug(f"📑 СОДЕРЖАНИЕ (прочее): {text_clean[:50]}...")
                return "skip"
        
        # 3. ВСЕГДА проверяем служебные разделы (титульная страница, задания и т.д.)
        if self._is_title_page_content(text_clean) or self._is_service_content(text_clean):
            logger.debug(f"🔴 СЛУЖЕБНЫЙ РАЗДЕЛ: {text_clean[:50]}...")
            return "skip"
        
        # 4. Проверяем, начинается ли основное содержание (если мы еще не в нем)
        if not self.document_state['found_main_content'] and self._is_main_content_start(text_clean):
            logger.info(f"🟢 НАЧАЛО ОСНОВНОГО СОДЕРЖАНИЯ: {text_clean[:60]}...")
            self.document_state['in_title_section'] = False
            self.document_state['found_main_content'] = True
            
            # Определяем тип этого параграфа (скорее всего H1)
            if self._is_h1_paragraph_content(text_clean):
                return "h1"
            elif self._is_h2_paragraph_content(text_clean):
                return "h2"
            else:
                return "regular"
        
        # 5. Если мы все еще в титульной секции
        if self.document_state['in_title_section']:
            # Если это не явный маркер, но мы еще не нашли основное содержание
            # продолжаем пропускать (может быть продолжение титульной страницы)
            logger.debug(f"⚪ ПРОПУСК (титульная): {text_clean[:50]}...")
            return "skip"
        
        # 4. Мы уже в основном содержании - классифицируем как обычно
        
        # Проверяем H1
        if self._is_h1_paragraph_content(text_clean):
            return "h1"
        
        # Проверяем H2
        if self._is_h2_paragraph_content(text_clean):
            return "h2"
        
        # Проверяем список
        if self._is_list_paragraph(text_clean):
            return "list"
        
        # Обычный параграф
        return "regular"
    
    def _is_main_content_start(self, text: str) -> bool:
        """Определяет начало основного содержания ВКР"""
        
        text_upper = text.upper().strip()
        
        # Исключаем строки содержания (с точками и номерами страниц)
        if self._is_contents_line(text):
            return False
        
        # Маркеры начала основного содержания (только точные совпадения!)
        main_content_markers = [
            "ВВЕДЕНИЕ",
            "ГЛАВА 1",
            "1. ВВЕДЕНИЕ",
            "1 ВВЕДЕНИЕ", 
            "CHAPTER 1",
            "РЕФЕРАТ",
            "ABSTRACT",
            "АННОТАЦИЯ"
        ]
        
        # Только ТОЧНЫЕ совпадения для безопасности
        for marker in main_content_markers:
            if text_upper == marker:
                return True
        
        # Паттерны для глав (но только если это НЕ содержит номер страницы в конце)
        chapter_patterns = [
            r"^ГЛАВА\s+\d+$",  # "ГЛАВА 1", "ГЛАВА 2" (но НЕ "ГЛАВА 1   10")
            r"^\d+\.\s*[А-ЯЁ][А-ЯЁа-яё\s]*$",  # "1. ВВЕДЕНИЕ" (но НЕ "1. ВВЕДЕНИЕ   8")
            r"^\d+\s+[А-ЯЁ][А-ЯЁа-яё\s]*$",    # "1 ВВЕДЕНИЕ" (но НЕ "1 ВВЕДЕНИЕ  8")
        ]
        
        for pattern in chapter_patterns:
            if re.match(pattern, text_upper):
                # Дополнительная проверка: не заканчивается ли на номер страницы
                if not re.search(r'\s+\d+$', text.strip()):
                    return True
        
        return False
    
    def _is_title_page_content(self, text: str) -> bool:
        """Определяет содержимое титульного листа"""
        
        text_upper = text.upper()
        
        # Строгие маркеры титульного листа
        title_markers = [
            "ДИПЛОМНАЯ РАБОТА",
            "ВЫПУСКНАЯ КВАЛИФИКАЦИОННАЯ РАБОТА", 
            "МИНИСТЕРСТВО ОБРАЗОВАНИЯ",
            "МИНИСТЕРСТВО НАУКИ",
            "ФЕДЕРАЛЬНОЕ ГОСУДАРСТВЕННОЕ",
            "ОБРАЗОВАТЕЛЬНОЕ УЧРЕЖДЕНИЕ",
            "ВЫСШЕГО ОБРАЗОВАНИЯ",
            "КАФЕДРА",
            "НАПРАВЛЕНИЕ ПОДГОТОВКИ",
            "ПРОФИЛЬ",
            "ТЕМА:",
            "ВЫПОЛНИЛ:",
            "СТУДЕНТ",
            "ГРУППЫ",
            "НАУЧНЫЙ РУКОВОДИТЕЛЬ",
            "КОНСУЛЬТАНТ",
            "ДОПУЩЕН К ЗАЩИТЕ",
            "РАБОТА ВЫПОЛНЕНА",
            "ОЦЕНКА",
            "ПОДПИСЬ"
        ]
        
        for marker in title_markers:
            if marker in text_upper:
                return True
        
        # Паттерны для ФИО и должностей
        fio_patterns = [
            r"[А-ЯЁ][а-яё]+\s+[А-ЯЁ]\.[А-ЯЁ]\.",  # Иванов И.И.
            r"[А-ЯЁ][а-яё]+\s+[А-ЯЁ][а-яё]+\s+[А-ЯЁ][а-яё]+",  # Иванов Иван Иванович
        ]
        
        for pattern in fio_patterns:
            if re.search(pattern, text):
                return True
        
        # Если текст короткий и состоит в основном из заглавных букв (заголовок титульника)
        if len(text) < 200:
            alpha_chars = [c for c in text if c.isalpha()]
            if alpha_chars:
                upper_ratio = sum(1 for c in alpha_chars if c.isupper()) / len(alpha_chars)
                # Для титульника характерны короткие строки с высоким процентом заглавных букв
                if upper_ratio > 0.8 and len(text.split()) <= 5:
                    return True
        
        return False
    
    def _is_contents_header(self, text: str) -> bool:
        """Определяет заголовок страницы содержания"""
        
        text_upper = text.upper().strip()
        
        # Заголовки содержания
        content_headers = [
            "СОДЕРЖАНИЕ",
            "ОГЛАВЛЕНИЕ", 
            "CONTENTS",
            "TABLE OF CONTENTS"
        ]
        
        # Точное совпадение для заголовков содержания
        for header in content_headers:
            if text_upper == header:
                return True
        
        return False
    
    def _is_contents_line(self, text: str) -> bool:
        """Определяет строку содержания (с точками, пробелами и номерами страниц)"""
        
        text_clean = text.strip()
        
        if not text_clean:
            return False
        
        # Паттерны для строк содержания
        content_patterns = [
            # С точками и номерами страниц
            r".+\.{3,}.+\d+$",  # "Введение...........3"
            r".+\.{2,}\s*\d+$",  # "1. Обзор литературы..5" 
            r"^[А-ЯЁ\d\.\s]+\.{3,}\d+$",  # "ГЛАВА 1...10"
            r"^\d+[\.\s][А-ЯЁа-яё\s]+\.{3,}\d+$",  # "1 Введение.....4"
            r"^\d+\.\d+[\.\s][А-ЯЁа-яё\s]+\.{3,}\d+$",  # "1.1 Подраздел.....8"
            
            # С пробелами и номерами страниц (без точек)
            r"^[А-ЯЁа-яё\s]+\s+\d+$",  # "Введение    8", "Заключение   50"
            r"^\d+\.\s*[А-ЯЁа-яё\s]+\s+\d+$",  # "1. Анализ предметной области   11"
            r"^\d+\s+[А-ЯЁа-яё\s]+\s+\d+$",  # "1 Введение   4"
            r"^\d+\.\d+\s+[А-ЯЁа-яё\s]+\s+\d+$",  # "1.1 Недостатки   11"
            
            # Специальные случаи
            r"^[А-ЯЁа-яё\s,]+\d+$",  # "Определения, обозначения и сокращения5"
        ]
        
        for pattern in content_patterns:
            if re.search(pattern, text_clean):
                return True
        
        # Дополнительная проверка: если строка заканчивается на число и содержит мало слов
        # (характерно для строк содержания)
        if re.search(r'\d+$', text_clean):
            words = text_clean.split()
            if len(words) >= 2 and len(words) <= 8:  # Обычно строки содержания содержат 2-8 слов
                # Последнее слово - число (номер страницы)
                try:
                    int(words[-1])
                    return True
                except ValueError:
                    pass
        
        return False
    
    def _is_service_content(self, text: str) -> bool:
        """Определяет другие служебные разделы (задание, календарный план, содержание и т.д.)"""
        
        text_upper = text.upper()
        text_clean = text.strip()
        
        # Основные служебные разделы
        service_markers = [
            "ЗАДАНИЕ НА",
            "КАЛЕНДАРНЫЙ ПЛАН",
            "КАЛЕНДАРНО-ТЕМАТИЧЕСКИЙ",
            "ТЕХНИЧЕСКОЕ ЗАДАНИЕ",
            "УТВЕРЖДАЮ",
            "РАССМОТРЕНО",
            "СОГЛАСОВАНО",
            "ОТЗЫВ",
            "РЕЦЕНЗИЯ",
            "СПРАВКА О ВНЕДРЕНИИ",
            "АКТ О ВНЕДРЕНИИ"
        ]
        
        for marker in service_markers:
            if marker in text_upper:
                return True
        

        
        return False
    
    def _is_h1_paragraph_content(self, text: str) -> bool:
        """Проверяет, является ли параграф заголовком H1 (только для основного содержания)"""
        
        patterns = self.requirements["h1_formatting"]["detection_patterns"]
        
        for pattern in patterns:
            if re.match(pattern, text.upper().strip()):
                return True
        
        # Дополнительная проверка: короткий текст с большим процентом заглавных букв
        if len(text) < 100:
            alpha_chars = [c for c in text if c.isalpha()]
            if alpha_chars:
                upper_ratio = sum(1 for c in alpha_chars if c.isupper()) / len(alpha_chars)
                if upper_ratio > 0.7:
                    return True
        
        return False
    
    def _is_h2_paragraph_content(self, text: str) -> bool:
        """Проверяет, является ли параграф заголовком H2 (только для основного содержания)"""
        
        patterns = self.requirements["h2_formatting"]["detection_patterns"]
        
        for pattern in patterns:
            if re.match(pattern, text.strip()):
                return True
        
        return False
    
    def _is_list_paragraph(self, text: str) -> bool:
        """Проверяет, является ли параграф элементом списка"""
        
        patterns = self.requirements["lists"]["bullet_lists"]["detection_patterns"]
        
        for pattern in patterns:
            if re.match(pattern, text):
                return True
        
        return False
    
    def _format_h1_paragraph(self, paragraph) -> None:
        """Форматирует заголовок H1"""
        
        try:
            h1_config = self.requirements["h1_formatting"]
            
            # Добавляем разрыв страницы если нужно
            if h1_config["page_break_before"] and self._not_first_paragraph(paragraph):
                self._add_page_break_before(paragraph)
            
            # Форматируем текст
            self._apply_font_formatting(paragraph, h1_config)
            
            # Заглавные буквы
            if h1_config["text_transform"] == "uppercase":
                self._make_text_uppercase(paragraph)
            
            # Выравнивание
            paragraph.alignment = self.align_map[h1_config["alignment"]]
            
            # Отступы
            pf = paragraph.paragraph_format
            pf.space_before = Pt(h1_config["space_before_pt"])
            pf.space_after = Pt(h1_config["space_after_pt"])
            
            logger.debug(f"H1 отформатирован: {paragraph.text[:30]}...")
            
        except Exception as e:
            logger.error(f"Ошибка форматирования H1: {e}")
            raise
    
    def _format_h2_paragraph(self, paragraph) -> None:
        """Форматирует заголовок H2"""
        
        try:
            h2_config = self.requirements["h2_formatting"]
            
            # Форматируем текст
            self._apply_font_formatting(paragraph, h2_config)
            
            # Выравнивание
            paragraph.alignment = self.align_map[h2_config["alignment"]]
            
            # Отступы
            pf = paragraph.paragraph_format
            pf.space_before = Pt(h2_config["space_before_pt"])
            pf.space_after = Pt(h2_config["space_after_pt"])
            pf.left_indent = Cm(h2_config.get("paragraph_indent_cm", 0))
            
            logger.debug(f"H2 отформатирован: {paragraph.text[:30]}...")
            
        except Exception as e:
            logger.error(f"Ошибка форматирования H2: {e}")
            raise
    
    def _format_list_paragraph(self, paragraph) -> None:
        """Форматирует элемент списка"""
        
        try:
            list_config = self.requirements["lists"]["bullet_lists"]
            font_config = list_config["font"]
            
            # Форматируем шрифт
            self._apply_font_formatting(paragraph, {
                "font_name": font_config["name"],
                "font_size": font_config["size"]
            })
            
            # Выравнивание
            paragraph.alignment = self.align_map[list_config["alignment"]]
            
            # Отступ
            pf = paragraph.paragraph_format
            pf.left_indent = Cm(list_config["indent_cm"])
            
            # Междустрочный интервал
            line_spacing = font_config["line_spacing"]
            if line_spacing in self.line_spacing_map:
                pf.line_spacing_rule = self.line_spacing_map[line_spacing]
            
            logger.debug(f"Список отформатирован: {paragraph.text[:30]}...")
            
        except Exception as e:
            logger.error(f"Ошибка форматирования списка: {e}")
            raise
    
    def _format_regular_paragraph(self, paragraph) -> None:
        """Форматирует обычный параграф"""
        
        try:
            if not paragraph.text.strip():
                return
            
            base_config = self.requirements["base_formatting"]
            
            # Форматируем шрифт
            self._apply_font_formatting(paragraph, base_config)
            
            # Выравнивание
            paragraph.alignment = self.align_map[base_config["text_alignment"]]
            
            # Отступы и интервалы
            pf = paragraph.paragraph_format
            pf.first_line_indent = Cm(base_config["paragraph_indent_cm"])
            
            # Междустрочный интервал
            line_spacing = base_config["line_spacing"]
            if line_spacing in self.line_spacing_map:
                pf.line_spacing_rule = self.line_spacing_map[line_spacing]
            
        except Exception as e:
            logger.error(f"Ошибка форматирования обычного параграфа: {e}")
            raise
    
    def _apply_font_formatting(self, paragraph, config: Dict[str, Any]) -> None:
        """Применяет форматирование шрифта к параграфу"""
        
        # Создаем run если его нет
        if not paragraph.runs:
            paragraph.add_run()
        
        # Применяем ко всем runs
        for run in paragraph.runs:
            font = run.font
            
            if "font_name" in config:
                font.name = config["font_name"]
            
            if "font_size" in config:
                font.size = Pt(config["font_size"])
            
            if config.get("font_weight") == "bold":
                font.bold = True
    
    def _make_text_uppercase(self, paragraph) -> None:
        """Преобразует текст параграфа в верхний регистр"""
        
        original_text = paragraph.text
        paragraph.clear()
        run = paragraph.add_run(original_text.upper())
        
        # Сохраняем базовое форматирование
        font = run.font
        font.name = self.requirements["h1_formatting"]["font_name"]
        font.size = Pt(self.requirements["h1_formatting"]["font_size"])
        if self.requirements["h1_formatting"]["font_weight"] == "bold":
            font.bold = True
    
    def _add_page_break_before(self, paragraph) -> None:
        """Добавляет разрыв страницы перед параграфом"""
        
        if paragraph.runs:
            # Вставляем разрыв в начало первого run
            first_run = paragraph.runs[0]
            first_run.add_break(WD_BREAK.PAGE)
        else:
            # Создаем новый run с разрывом
            run = paragraph.add_run()
            run.add_break(WD_BREAK.PAGE)
    
    def _not_first_paragraph(self, target_paragraph) -> bool:
        """Проверяет, что параграф не первый в документе"""
        
        try:
            # Получаем документ
            doc = target_paragraph._parent
            while hasattr(doc, '_parent') and doc._parent is not None:
                doc = doc._parent
            
            # Проверяем позицию
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph == target_paragraph:
                    # Есть ли непустые параграфы до этого?
                    for j in range(i):
                        if doc.paragraphs[j].text.strip():
                            return True
                    return False
            
            return True
            
        except Exception:
            return True  # В случае ошибки считаем, что не первый
    
    def get_statistics(self) -> Dict[str, int]:
        """Возвращает статистику обработки"""
        stats = self.stats.copy()
        stats.update({
            'title_pages_detected': 1 if self.document_state['found_main_content'] else 0,
            'main_content_found': self.document_state['found_main_content'],
            'contents_section_detected': not self.document_state['in_contents_section'] and self.document_state['found_main_content']
        })
        return stats


# Главная функция для использования в API
def format_vkr_document(input_path: str, requirements: Dict[str, Any], output_path: str) -> tuple[bool, Dict[str, int]]:
    """
    Форматирует ВКР согласно требованиям
    
    Args:
        input_path: путь к исходному файлу ВКР
        requirements: словарь требований (из заглушки)
        output_path: путь к результирующему файлу
        
    Returns:
        tuple: (успех, статистика)
    """
    
    formatter = SimpleVKRFormatter(requirements)
    success = formatter.format_document(input_path, output_path)
    stats = formatter.get_statistics()
    
    return success, stats
