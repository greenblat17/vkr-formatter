from docx import Document
from docx.shared import Cm, Pt
from typing import Dict, Any, Tuple
from pathlib import Path

from paragraph_classifier import ParagraphClassifier
from style_based_classifier import StyleBasedClassifier
from paragraph_formatter import ParagraphFormatter
from statistics_tracker import StatisticsTracker
from document_state import logger


class VKRFormatter:
    """Основной класс для форматирования ВКР"""

    def __init__(self, requirements: Dict[str, Any], use_style_based_classification: bool = True, strict_style_mode: bool = False):
        self.requirements = requirements
        self.use_style_based = use_style_based_classification
        
        if use_style_based_classification:
            self.classifier = StyleBasedClassifier(requirements, strict_style_mode=strict_style_mode)
            if strict_style_mode:
                logger.info("🔒 Используем СТРОГУЮ классификацию на основе стилей (без fallback)")
            else:
                logger.info("🎨 Используем классификацию на основе стилей документа (с fallback)")
        else:
            self.classifier = ParagraphClassifier(requirements)
            logger.info("📝 Используем классификацию на основе текстовых паттернов")
            
        self.formatter = ParagraphFormatter(requirements)
        self.stats = StatisticsTracker()
        self.h1_count = 0  # Счетчик H1 заголовков (только для разрывов страниц)

    def format_document(self, input_path: str, output_path: str) -> bool:
        """Форматирует документ"""
        try:
            logger.info(f"📂 Начинаем форматирование: {input_path}")
            logger.info(f"💾 Выходной путь: {output_path}")

            # Проверяем входной файл
            input_file = Path(input_path)
            if not input_file.exists():
                logger.error(f"❌ Входной файл не существует: {input_path}")
                return False

            # Загружаем документ
            logger.info("📖 Загружаем документ...")
            doc = Document(input_path)
            logger.info(
                f"✅ Документ загружен, параграфов: {len(doc.paragraphs)}")

            # Применяем глобальные настройки
            logger.info("⚙️  Применяем глобальные настройки...")
            self._apply_global_settings(doc)

            # Обрабатываем параграфы
            logger.info("🔄 Обрабатываем параграфы...")
            self._process_all_paragraphs(doc)
            
            # Обрабатываем таблицы
            logger.info("📊 Обрабатываем таблицы...")
            self._process_all_tables(doc)

            # Сохраняем результат
            logger.info(f"💾 Сохраняем документ в: {output_path}")
            doc.save(output_path)

            # Проверяем результат
            output_file = Path(output_path)
            if output_file.exists():
                logger.info(
                    f"✅ Файл создан, размер: {output_file.stat().st_size} байт")
            else:
                logger.error(f"❌ Файл НЕ создался: {output_path}")
                return False

            final_stats = self.get_statistics()
            logger.info(
                f"🎉 Форматирование завершено! Статистика: {final_stats}")
            return True

        except Exception as e:
            logger.error(f"Ошибка форматирования: {e}")
            import traceback
            logger.error(f"Полная трассировка: {traceback.format_exc()}")
            return False

    def _apply_global_settings(self, doc: Document) -> None:
        """Применяет глобальные настройки документа"""
        try:
            base_config = self.requirements["base_formatting"]
            
            # Применяем поля страницы
            margins = base_config["margins_cm"]
            for section in doc.sections:
                section.top_margin = Cm(margins["top"])
                section.bottom_margin = Cm(margins["bottom"])
                section.left_margin = Cm(margins["left"])
                section.right_margin = Cm(margins["right"])
            
            logger.info(f"✅ Применены поля: {margins}")
            
            # Настраиваем стили документа по умолчанию
            self._configure_default_styles(doc, base_config)
            
            logger.info("✅ Глобальные настройки применены")

        except Exception as e:
            logger.error(f"Ошибка применения глобальных настроек: {e}")
            self.stats.increment('errors')
    

    
    def _configure_default_styles(self, doc: Document, base_config: Dict[str, Any]) -> None:
        """Настраивает стили документа по умолчанию (только базовые настройки)"""
        try:
            # Получаем стиль Normal (базовый стиль)
            styles = doc.styles
            normal_style = styles['Normal']
            
            # Настраиваем только базовый шрифт
            font = normal_style.font
            font.name = base_config["font_name"]
            font.size = Pt(base_config["font_size"])
            
            logger.info("✅ Стиль Normal настроен (только шрифт)")
            
        except Exception as e:
            logger.warning(f"Не удалось настроить стили по умолчанию: {e}")
    


    def _process_all_paragraphs(self, doc: Document) -> None:
        """Обрабатывает все параграфы документа"""
        logger.info("Начинаем обработку параграфов...")

        for i, paragraph in enumerate(doc.paragraphs):
            self.stats.increment('total_paragraphs')

            try:
                text = paragraph.text.strip()
                
                # Выбираем метод классификации
                if self.use_style_based:
                    paragraph_type = self.classifier.classify_paragraph_by_style(paragraph, text)
                else:
                    paragraph_type = self.classifier.classify_paragraph(text)

                # Логируем непустые параграфы
                if text:
                    logger.debug(
                        f"Параграф {i+1}: тип='{paragraph_type}', текст='{text[:100]}{'...' if len(text) > 100 else ''}'")

                # Применяем форматирование
                self._apply_paragraph_formatting(
                    paragraph, paragraph_type, i+1, text)

            except Exception as e:
                logger.warning(f"Ошибка обработки параграфа {i+1}: {e}")
                self.stats.increment('errors')

        final_stats = self.stats.stats
        logger.info(
            f"Обработка параграфов завершена. Статистика: {final_stats}")

    def _process_all_tables(self, doc: Document) -> None:
        """Обрабатывает все таблицы в документе"""
        logger.info("Начинаем обработку таблиц...")
        
        table_count = 0
        
        try:
            # Получаем все таблицы из документа
            tables = doc.tables
            
            if not tables:
                logger.info("📊 Таблицы в документе не найдены")
                return
            
            logger.info(f"📊 Найдено таблиц: {len(tables)}")
            
            for i, table in enumerate(tables):
                table_count += 1
                
                try:
                    logger.info(f"📊 Обрабатываем таблицу #{i+1}")
                    
                    # Форматируем таблицу
                    self.formatter.format_table(table)
                    self.stats.increment('tables_formatted')
                    
                    logger.info(f"✅ Таблица #{i+1} отформатирована")
                    
                except Exception as e:
                    logger.warning(f"❌ Ошибка обработки таблицы #{i+1}: {e}")
                    self.stats.increment('errors')
            
            logger.info(f"📊 Обработка таблиц завершена. Обработано: {table_count}")
            
        except Exception as e:
            logger.error(f"❌ Ошибка при обработке таблиц: {e}")
            self.stats.increment('errors')

    def _apply_paragraph_formatting(self, paragraph, paragraph_type: str, index: int, text: str) -> None:
        """Применяет форматирование к параграфу"""
        if paragraph_type == "skip":
            self.stats.increment('skipped_paragraphs')
            logger.info(
                f"⏭️  ПРОПУСК #{index}: {text[:60]}{'...' if len(text) > 60 else ''}")

        elif paragraph_type == "h1":
            # Увеличиваем счетчик H1 (для разрывов страниц)
            self.h1_count += 1
            
            # Передаем счетчик H1 в форматтер (для разрывов страниц)
            self.formatter.format_h1(paragraph, self.h1_count - 1)
            
            self.stats.increment('h1_formatted')
            logger.info(f"📝 H1 #{index}: {text[:40]}...")

        elif paragraph_type == "h2":
            self.formatter.format_h2(paragraph)
            self.stats.increment('h2_formatted')
            logger.info(f"📄 H2 #{index}: {text[:40]}...")

        elif paragraph_type == "h3":
            self.formatter.format_h3(paragraph)
            self.stats.increment('h3_formatted')
            logger.info(f"📋 H3 #{index}: {text[:40]}...")

        elif paragraph_type == "h4":
            self.formatter.format_h4(paragraph)
            self.stats.increment('h4_formatted')
            logger.info(f"📌 H4 #{index}: {text[:40]}...")

        elif paragraph_type == "list":
            self.formatter.format_list(paragraph)
            self.stats.increment('lists_formatted')
            logger.debug(f"📋 СПИСОК #{index}: {text[:40]}...")

        elif paragraph_type == "references_header":
            self.formatter.format_references_header(paragraph)
            self.stats.increment('references_headers_formatted')
            logger.info(f"📚 ЗАГОЛОВОК СПИСКА ЛИТЕРАТУРЫ #{index}: {text[:40]}...")

        elif paragraph_type == "bibliography_entry":
            self.formatter.format_bibliography_entry(paragraph)
            self.stats.increment('bibliography_entries_formatted')
            logger.info(f"📖 БИБЛИОГРАФИЧЕСКАЯ ЗАПИСЬ #{index}: {text[:60]}...")

        elif paragraph_type == "bibliography_continuation":
            self.formatter.format_bibliography_continuation(paragraph)
            self.stats.increment('bibliography_continuations_formatted')
            logger.info(f"📄 ПРОДОЛЖЕНИЕ ЗАПИСИ #{index}: {text[:60]}...")

        elif paragraph_type == "references_text":
            self.formatter.format_references_text(paragraph)
            self.stats.increment('references_text_formatted')
            logger.debug(f"📝 ТЕКСТ В СПИСКЕ ЛИТЕРАТУРЫ #{index}: {text[:40]}...")

        elif paragraph_type.startswith("special_"):
            section_name = paragraph_type.replace("special_", "")
            self.formatter.format_special_section(paragraph, section_name)
            self.stats.increment(f'special_{section_name}_formatted')
            logger.info(f"⭐ СПЕЦИАЛЬНЫЙ РАЗДЕЛ ({section_name.upper()}) #{index}: {text[:40]}...")

        elif paragraph_type == "table_caption":
            self.formatter.format_table_caption(paragraph)
            self.stats.increment('table_captions_formatted')
            logger.debug(f"📊 ПОДПИСЬ ТАБЛИЦЫ #{index}: {text[:40]}...")

        elif paragraph_type == "figure_image":
            self.formatter.format_figure_image(paragraph)
            self.stats.increment('figure_images_formatted')
            logger.info(f"🖼️ ИЗОБРАЖЕНИЕ РИСУНКА #{index}")

        elif paragraph_type == "figure_caption":
            self.formatter.format_figure_caption(paragraph)
            self.stats.increment('figure_captions_formatted')
            logger.info(f"🖼️ ПОДПИСЬ РИСУНКА #{index}: {text[:60]}...")

        elif paragraph_type == "formula":
            self.formatter.format_formula(paragraph)
            self.stats.increment('formulas_formatted')
            logger.info(f"🔢 ФОРМУЛА #{index}: {text[:40]}...")

        elif paragraph_type == "formula_numbering":
            self.formatter.format_formula_numbering(paragraph)
            self.stats.increment('formula_numbering_formatted')
            logger.info(f"🔢 НУМЕРАЦИЯ ФОРМУЛЫ #{index}: {text[:40]}...")

        elif paragraph_type == "formula_explanation":
            self.formatter.format_formula_explanation(paragraph)
            self.stats.increment('formula_explanations_formatted')
            logger.info(f"🔤 ПОЯСНЕНИЕ К ФОРМУЛЕ #{index}: {text[:60]}...")

        else:  # regular
            self.formatter.format_regular(paragraph)
            self.stats.increment('regular_formatted')
            # Дополнительное логирование для отладки
            if self.classifier.get_state().in_references_section:
                logger.warning(f"⚠️  ВНИМАНИЕ: Строка в списке литературы классифицирована как 'regular': {text[:60]}...")

    def get_statistics(self) -> Dict[str, Any]:
        """Возвращает статистику обработки"""
        return self.stats.get_statistics(self.classifier.get_state())


def format_vkr_document(input_path: str, requirements: Dict[str, Any], output_path: str, use_style_based: bool = True, strict_style_mode: bool = False) -> Tuple[bool, Dict[str, Any]]:
    """
    Форматирует ВКР согласно требованиям

    Args:
        input_path: путь к исходному файлу ВКР
        requirements: словарь требований
        output_path: путь к результирующему файлу
        use_style_based: использовать классификацию на основе стилей (по умолчанию True)
        strict_style_mode: строгий режим стилей - игнорировать паттерны для Normal стиля

    Returns:
        tuple: (успех, статистика)
    """
    formatter = VKRFormatter(requirements, use_style_based_classification=use_style_based, strict_style_mode=strict_style_mode)
    success = formatter.format_document(input_path, output_path)
    stats = formatter.get_statistics()

    return success, stats
