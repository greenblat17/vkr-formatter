from docx.shared import Pt, Cm
from docx.enum.text import WD_BREAK
from typing import Dict, Any
from formatting_constants import FormattingConstants
from document_state import logger


class ParagraphFormatter:
    """Форматирует параграфы разных типов"""

    def __init__(self, requirements: Dict[str, Any]):
        self.requirements = requirements

    def format_h1(self, paragraph, h1_count_before: int = 0) -> None:
        """Форматирует заголовок H1 согласно требованиям ГОСТ"""
        try:
            config = self.requirements["h1_formatting"]
            
            logger.info(f"🔤 Форматирование H1: {paragraph.text[:50]}...")

            # 1. Сначала преобразуем в заглавные буквы (если нужно)
            if config.get("text_transform") == "uppercase":
                self._make_text_uppercase(paragraph, config)
                logger.debug("   ↳ Текст преобразован в ЗАГЛАВНЫЕ БУКВЫ")

            # 2. Применяем шрифт и размер
            self._apply_font_formatting(paragraph, config)
            logger.debug(f"   ↳ Шрифт: {config['font_name']}, {config['font_size']}pt, жирный")

            # 3. Разрыв страницы перед заголовком (кроме самого первого H1 в документе)
            # ВАЖНО: добавляем разрыв ПОСЛЕ всех текстовых преобразований!
            if config.get("page_break_before", False) and h1_count_before > 0:
                self._add_page_break_before(paragraph)
                logger.debug(f"   ↳ Добавлен разрыв страницы (H1 #{h1_count_before + 1})")
            else:
                logger.debug(f"   ↳ Разрыв страницы НЕ добавлен (первый H1 в документе)")

            # 4. Выравнивание по центру
            paragraph.alignment = FormattingConstants.ALIGN_MAP[config["alignment"]]
            logger.debug(f"   ↳ Выравнивание: {config['alignment']}")

            # 5. Настройки параграфа
            pf = paragraph.paragraph_format
            
            # Отступы до и после заголовка
            pf.space_before = Pt(config["space_before_pt"])
            pf.space_after = Pt(config["space_after_pt"])
            
            # Убираем отступ первой строки для заголовков
            pf.first_line_indent = Cm(0)
            pf.left_indent = Cm(0)
            pf.right_indent = Cm(0)
            
            # Междустрочный интервал для заголовка (обычно одинарный)
            pf.line_spacing_rule = FormattingConstants.LINE_SPACING_MAP.get(1.0)
            
            logger.debug(f"   ↳ Отступы: до={config['space_before_pt']}pt, после={config['space_after_pt']}pt")
            logger.info(f"✅ H1 отформатирован: {paragraph.text[:40]}...")

        except Exception as e:
            logger.error(f"❌ Ошибка форматирования H1: {e}")
            raise

    def format_h2(self, paragraph) -> None:
        """Форматирует заголовок H2"""
        try:
            config = self.requirements["h2_formatting"]

            self._apply_font_formatting(paragraph, config)
            paragraph.alignment = FormattingConstants.ALIGN_MAP[config["alignment"]]

            pf = paragraph.paragraph_format
            pf.space_before = Pt(config["space_before_pt"])
            pf.space_after = Pt(config["space_after_pt"])
            pf.left_indent = Cm(config.get("paragraph_indent_cm", 0))

            logger.debug(f"H2 отформатирован: {paragraph.text[:30]}...")

        except Exception as e:
            logger.error(f"Ошибка форматирования H2: {e}")
            raise

    def format_h3(self, paragraph) -> None:
        """Форматирует заголовок H3"""
        try:
            config = self.requirements["h3_formatting"]

            self._apply_font_formatting(paragraph, config)
            paragraph.alignment = FormattingConstants.ALIGN_MAP[config["alignment"]]

            pf = paragraph.paragraph_format
            pf.space_before = Pt(config["space_before_pt"])
            pf.space_after = Pt(config["space_after_pt"])
            pf.left_indent = Cm(config.get("paragraph_indent_cm", 0))

            logger.debug(f"H3 отформатирован: {paragraph.text[:30]}...")

        except Exception as e:
            logger.error(f"Ошибка форматирования H3: {e}")
            raise

    def format_h4(self, paragraph) -> None:
        """Форматирует заголовок H4"""
        try:
            config = self.requirements["h4_formatting"]

            self._apply_font_formatting(paragraph, config)
            paragraph.alignment = FormattingConstants.ALIGN_MAP[config["alignment"]]

            pf = paragraph.paragraph_format
            pf.space_before = Pt(config["space_before_pt"])
            pf.space_after = Pt(config["space_after_pt"])
            pf.left_indent = Cm(config.get("paragraph_indent_cm", 0))

            logger.debug(f"H4 отформатирован: {paragraph.text[:30]}...")

        except Exception as e:
            logger.error(f"Ошибка форматирования H4: {e}")
            raise

    def format_list(self, paragraph) -> None:
        """Форматирует элемент списка"""
        try:
            config = self.requirements["lists"]["bullet_lists"]
            font_config = config["font"]

            self._apply_font_formatting(paragraph, {
                "font_name": font_config["name"],
                "font_size": font_config["size"]
            })

            paragraph.alignment = FormattingConstants.ALIGN_MAP[config["alignment"]]

            pf = paragraph.paragraph_format
            pf.left_indent = Cm(config["indent_cm"])

            line_spacing = font_config["line_spacing"]
            if line_spacing in FormattingConstants.LINE_SPACING_MAP:
                pf.line_spacing_rule = FormattingConstants.LINE_SPACING_MAP[line_spacing]

            logger.debug(f"Список отформатирован: {paragraph.text[:30]}...")

        except Exception as e:
            logger.error(f"Ошибка форматирования списка: {e}")
            raise

    def format_regular(self, paragraph) -> None:
        """Форматирует обычный параграф согласно базовым требованиям ГОСТ"""
        try:
            if not paragraph.text.strip():
                return

            config = self.requirements["base_formatting"]

            # Применяем шрифт и размер
            self._apply_font_formatting(paragraph, config)
            
            # Выравнивание по ширине
            paragraph.alignment = FormattingConstants.ALIGN_MAP[config["text_alignment"]]

            # Настройки параграфа
            pf = paragraph.paragraph_format
            
            # Отступ первой строки (красная строка)
            pf.first_line_indent = Cm(config["paragraph_indent_cm"])
            
            # Междустрочный интервал
            line_spacing = config["line_spacing"]
            if line_spacing in FormattingConstants.LINE_SPACING_MAP:
                pf.line_spacing_rule = FormattingConstants.LINE_SPACING_MAP[line_spacing]
            
            # Убираем дополнительные отступы между параграфами
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            
            # Убираем левый отступ (только красная строка)
            pf.left_indent = Cm(0)
            pf.right_indent = Cm(0)

            logger.debug(f"Обычный параграф отформатирован: {paragraph.text[:30]}...")

        except Exception as e:
            logger.error(f"Ошибка форматирования обычного параграфа: {e}")
            raise

    def _apply_font_formatting(self, paragraph, config: Dict[str, Any]) -> None:
        """Применяет форматирование шрифта"""
        if not paragraph.runs:
            paragraph.add_run()

        for run in paragraph.runs:
            font = run.font

            if "font_name" in config:
                font.name = config["font_name"]

            if "font_size" in config:
                font.size = Pt(config["font_size"])

            if config.get("font_weight") == "bold":
                font.bold = True

    def _make_text_uppercase(self, paragraph, config: Dict[str, Any]) -> None:
        """Преобразует текст в верхний регистр"""
        original_text = paragraph.text
        paragraph.clear()
        run = paragraph.add_run(original_text.upper())

        font = run.font
        font.name = config["font_name"]
        font.size = Pt(config["font_size"])
        if config["font_weight"] == "bold":
            font.bold = True

    def _add_page_break_before(self, paragraph) -> None:
        """Добавляет разрыв страницы перед параграфом используя свойства параграфа"""
        try:
            # Используем свойство page_break_before параграфа
            # Это более элегантное решение, чем добавление элементов разрыва
            pf = paragraph.paragraph_format
            pf.page_break_before = True
            
            logger.debug(f"   ✅ Разрыв страницы установлен для параграфа: {paragraph.text[:30]}...")
            
        except Exception as e:
            logger.error(f"   ❌ Ошибка установки разрыва страницы: {e}")
            # Fallback: используем старый метод с элементами разрыва
            try:
                if paragraph.runs:
                    first_run = paragraph.runs[0]
                    # Сохраняем текст первого run
                    original_text = first_run.text
                    # Очищаем run
                    first_run.clear()
                    # Добавляем разрыв страницы
                    first_run.add_break(WD_BREAK.PAGE)
                    # Возвращаем текст
                    first_run.add_text(original_text)
                else:
                    # Если нет runs, создаем новый с разрывом
                    run = paragraph.add_run()
                    run.add_break(WD_BREAK.PAGE)
                    
                logger.debug(f"   ⚠️  Использован fallback метод для разрыва страницы")
                
            except Exception as fallback_error:
                logger.error(f"   ❌ Fallback метод также не сработал: {fallback_error}")

    def _should_add_page_break_for_h1(self, target_paragraph) -> bool:
        """Определяет, нужен ли разрыв страницы для H1 заголовка"""
        try:
            # Используем счетчик H1 заголовков, который уже был отформатирован
            # Если это первый H1 в процессе форматирования - не добавляем разрыв
            # Если это второй и последующие - добавляем разрыв
            
            # Получаем документ
            doc = target_paragraph._parent
            while hasattr(doc, '_parent') and doc._parent is not None:
                doc = doc._parent

            # Считаем количество H1 заголовков ДО текущего
            h1_count_before = 0
            target_found = False
            
            for paragraph in doc.paragraphs:
                if paragraph == target_paragraph:
                    target_found = True
                    break
                    
                # Проверяем, является ли параграф H1 заголовком
                if self._is_h1_heading(paragraph):
                    h1_count_before += 1
            
            # Если целевой параграф не найден, это ошибка
            if not target_found:
                logger.warning("   ⚠️  Целевой параграф не найден в документе")
                return False
            
            # Простая логика: если это НЕ первый H1 в документе - добавляем разрыв
            should_break = h1_count_before > 0
            
            logger.debug(f"   🔍 H1 заголовков до текущего: {h1_count_before}, разрыв страницы: {should_break}")
            return should_break

        except Exception as e:
            logger.warning(f"Ошибка определения разрыва страницы для H1: {e}")
            # В случае ошибки НЕ добавляем разрыв (безопасный вариант)
            return False

    def _is_h1_heading(self, paragraph) -> bool:
        """Проверяет, является ли параграф H1 заголовком"""
        try:
            # Проверяем по стилю
            if hasattr(paragraph, 'style') and paragraph.style:
                style_name = paragraph.style.name
                h1_styles = [
                    "Heading 1", "Заголовок 1", "Title", "Название", "Header 1", "H1"
                ]
                
                # Точное совпадение
                if style_name in h1_styles:
                    return True
                
                # Частичное совпадение
                style_lower = style_name.lower()
                for h1_style in h1_styles:
                    if h1_style.lower() in style_lower:
                        return True
            
            # Дополнительная проверка по тексту (fallback)
            text = paragraph.text.strip().upper()
            if text:
                # Простые паттерны для H1
                import re
                h1_patterns = [
                    r"^\d+\.\s*[А-ЯЁ\s]+$",           # "1. ВВЕДЕНИЕ"
                    r"^ГЛАВА\s+\d+",                   # "ГЛАВА 1"
                    r"^(ВВЕДЕНИЕ|ЗАКЛЮЧЕНИЕ|РЕФЕРАТ)$",  # специальные разделы
                    r"^[IVX]+\.\s*[А-ЯЁ\s]+$"        # "I. ВВЕДЕНИЕ"
                ]
                
                for pattern in h1_patterns:
                    if re.match(pattern, text):
                        return True
            
            return False
            
        except Exception:
            return False

    def _not_first_paragraph(self, target_paragraph) -> bool:
        """Проверяет, что параграф не первый (устаревший метод)"""
        try:
            doc = target_paragraph._parent
            while hasattr(doc, '_parent') and doc._parent is not None:
                doc = doc._parent

            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph == target_paragraph:
                    for j in range(i):
                        if doc.paragraphs[j].text.strip():
                            return True
                    return False

            return True

        except Exception:
            return True

    def format_references_header(self, paragraph) -> None:
        """Форматирует заголовок списка использованных источников"""
        try:
            # Используем форматирование H1 для заголовка списка литературы
            h1_config = self.requirements["h1_formatting"]
            
            logger.info(f"📚 Форматирование заголовка списка литературы: {paragraph.text[:50]}...")

            # Применяем шрифт и размер
            self._apply_font_formatting(paragraph, h1_config)
            
            # Преобразуем в заглавные буквы
            if h1_config.get("text_transform") == "uppercase":
                self._make_text_uppercase(paragraph, h1_config)
            
            # Выравнивание по центру
            paragraph.alignment = FormattingConstants.ALIGN_MAP[h1_config["alignment"]]
            
            # Настройки параграфа
            pf = paragraph.paragraph_format
            pf.space_before = Pt(h1_config["space_before_pt"])
            pf.space_after = Pt(h1_config["space_after_pt"])
            pf.first_line_indent = Cm(0)
            pf.left_indent = Cm(0)
            pf.right_indent = Cm(0)
            pf.line_spacing_rule = FormattingConstants.LINE_SPACING_MAP.get(1.0)
            
            # Добавляем разрыв страницы перед списком литературы
            if h1_config.get("page_break_before", False):
                self._add_page_break_before(paragraph)
                logger.debug("   ↳ Добавлен разрыв страницы перед списком литературы")
            
            logger.info(f"✅ Заголовок списка литературы отформатирован")

        except Exception as e:
            logger.error(f"❌ Ошибка форматирования заголовка списка литературы: {e}")
            raise

    def format_bibliography_entry(self, paragraph) -> None:
        """Форматирует библиографическую запись"""
        try:
            config = self.requirements["special_sections"]["references"]["content"]
            
            logger.debug(f"📖 Форматирование библиографической записи: {paragraph.text[:60]}...")

            # Применяем шрифт
            self._apply_font_formatting(paragraph, config)
            
            # Выравнивание по ширине
            paragraph.alignment = FormattingConstants.ALIGN_MAP[config["alignment"]]
            
            # Настройки параграфа
            pf = paragraph.paragraph_format
            
            # Обычный отступ первой строки (красная строка) для библиографических записей
            pf.first_line_indent = Cm(config["paragraph_indent_cm"])
            pf.left_indent = Cm(0)
            logger.debug(f"   ↳ Красная строка: first_line_indent={config['paragraph_indent_cm']}см")
            
            # Междустрочный интервал
            line_spacing = config["line_spacing"]
            if line_spacing in FormattingConstants.LINE_SPACING_MAP:
                pf.line_spacing_rule = FormattingConstants.LINE_SPACING_MAP[line_spacing]
            
            # Минимальные отступы между записями (согласно ГОСТ)
            pf.space_before = Pt(config.get("space_before_pt", 0))
            pf.space_after = Pt(config.get("space_after_pt", 0))
            
            logger.debug(f"✅ Библиографическая запись отформатирована")

        except Exception as e:
            logger.error(f"❌ Ошибка форматирования библиографической записи: {e}")
            raise

    def format_bibliography_continuation(self, paragraph) -> None:
        """Форматирует продолжение библиографической записи (без красной строки)"""
        try:
            config = self.requirements["special_sections"]["references"]["content"]
            
            logger.debug(f"📖 Форматирование продолжения библиографической записи: {paragraph.text[:60]}...")

            # Применяем шрифт
            self._apply_font_formatting(paragraph, config)
            
            # Выравнивание по ширине
            paragraph.alignment = FormattingConstants.ALIGN_MAP[config["alignment"]]
            
            # Настройки параграфа
            pf = paragraph.paragraph_format
            
            # БЕЗ красной строки для продолжения записи
            pf.first_line_indent = Cm(0)
            pf.left_indent = Cm(0)
            logger.debug(f"   ↳ Продолжение записи: БЕЗ красной строки")
            
            # Междустрочный интервал
            line_spacing = config["line_spacing"]
            if line_spacing in FormattingConstants.LINE_SPACING_MAP:
                pf.line_spacing_rule = FormattingConstants.LINE_SPACING_MAP[line_spacing]
            
            # Минимальные отступы между записями (согласно ГОСТ)
            pf.space_before = Pt(config.get("space_before_pt", 0))
            pf.space_after = Pt(config.get("space_after_pt", 0))
            
            logger.debug(f"✅ Продолжение библиографической записи отформатировано")

        except Exception as e:
            logger.error(f"❌ Ошибка форматирования продолжения библиографической записи: {e}")
            raise

    def format_references_text(self, paragraph) -> None:
        """Форматирует обычный текст в разделе списка литературы (теперь использует то же форматирование, что и bibliography_entry)"""
        try:
            # Используем то же форматирование, что и для библиографических записей
            self.format_bibliography_entry(paragraph)
            logger.debug(f"Текст в списке литературы отформатирован как библиографическая запись: {paragraph.text[:30]}...")

        except Exception as e:
            logger.error(f"Ошибка форматирования текста в списке литературы: {e}")
            raise

    def format_special_section(self, paragraph, section_name: str) -> None:
        """Форматирует специальные разделы (реферат, аннотация, введение, заключение, список литературы)"""
        try:
            if section_name in self.requirements["special_sections"]:
                config = self.requirements["special_sections"][section_name]
                
                logger.debug(f"⭐ Форматирование специального раздела '{section_name}': {paragraph.text[:40]}...")
                
                # Специальная обработка для заголовка списка литературы
                if section_name == "references":
                    logger.info(f"📚 Форматирование заголовка списка литературы: {paragraph.text[:50]}...")
                    
                    # Используем настройки для заголовка
                    title_config = config["title"]
                    
                    # Применяем шрифт и размер
                    self._apply_font_formatting(paragraph, title_config)
                    
                    # Преобразуем в заглавные буквы
                    if title_config.get("text_transform") == "uppercase":
                        self._make_text_uppercase(paragraph, title_config)
                    
                    # Выравнивание по центру
                    paragraph.alignment = FormattingConstants.ALIGN_MAP[title_config["alignment"]]
                    
                    # Настройки параграфа
                    pf = paragraph.paragraph_format
                    pf.space_before = Pt(title_config["space_before_pt"])
                    pf.space_after = Pt(title_config["space_after_pt"])
                    pf.first_line_indent = Cm(0)
                    pf.left_indent = Cm(0)
                    pf.right_indent = Cm(0)
                    pf.line_spacing_rule = FormattingConstants.LINE_SPACING_MAP.get(1.0)
                    
                    # Добавляем разрыв страницы перед списком литературы
                    if title_config.get("page_break_before", False):
                        self._add_page_break_before(paragraph)
                        logger.debug("   ↳ Добавлен разрыв страницы перед списком литературы")
                    
                    logger.info(f"✅ Заголовок списка литературы отформатирован")
                else:
                    # Обычное форматирование для других специальных разделов
                    self._apply_font_formatting(paragraph, config)
                    paragraph.alignment = FormattingConstants.ALIGN_MAP[config["alignment"]]
                    
                    pf = paragraph.paragraph_format
                    pf.first_line_indent = Cm(config["paragraph_indent_cm"])
                    
                    line_spacing = config["line_spacing"]
                    if line_spacing in FormattingConstants.LINE_SPACING_MAP:
                        pf.line_spacing_rule = FormattingConstants.LINE_SPACING_MAP[line_spacing]
                    
                    pf.space_before = Pt(0)
                    pf.space_after = Pt(0)
                    
                    logger.debug(f"✅ Специальный раздел '{section_name}' отформатирован")
            else:
                # Fallback к обычному форматированию
                self.format_regular(paragraph)
                logger.debug(f"⚠️ Неизвестный специальный раздел '{section_name}', применено обычное форматирование")

        except Exception as e:
            logger.error(f"❌ Ошибка форматирования специального раздела '{section_name}': {e}")
            raise

    def format_table_caption(self, paragraph) -> None:
        """Форматирует подпись таблицы согласно ГОСТ"""
        try:
            config = self.requirements["tables"]["caption"]
            
            text = paragraph.text.strip()
            logger.info(f"📊 Форматирование подписи таблицы: {text}")
            
            # Применяем форматирование шрифта
            self._apply_font_formatting(paragraph, config)
            logger.debug(f"   ↳ Шрифт: {config['font_name']} {config['font_size']}pt")
            
            # Выравнивание по левому краю (согласно ГОСТ)
            paragraph.alignment = FormattingConstants.ALIGN_MAP[config["alignment"]]
            logger.debug(f"   ↳ Выравнивание: {config['alignment']}")
            
            # Настройки параграфа
            pf = paragraph.paragraph_format
            
            # Отступы до и после подписи
            spacing_config = config["spacing"]
            pf.space_before = Pt(spacing_config["before_pt"])
            pf.space_after = Pt(spacing_config["after_pt"])
            logger.debug(f"   ↳ Отступы: до={spacing_config['before_pt']}pt, после={spacing_config['after_pt']}pt")
            
            # Убираем отступ первой строки для подписей
            pf.first_line_indent = Cm(0)
            pf.left_indent = Cm(0)
            pf.right_indent = Cm(0)
            
            # Междустрочный интервал
            line_spacing = config.get("line_spacing", 1.0)
            if line_spacing in FormattingConstants.LINE_SPACING_MAP:
                pf.line_spacing_rule = FormattingConstants.LINE_SPACING_MAP[line_spacing]
                logger.debug(f"   ↳ Междустрочный интервал: {line_spacing}")
            
            logger.info(f"✅ Подпись таблицы отформатирована: {paragraph.text[:50]}...")

        except Exception as e:
            logger.error(f"❌ Ошибка форматирования подписи таблицы: {e}")
            raise

    def format_table_content(self, paragraph) -> None:
        """Форматирует содержимое таблицы согласно ГОСТ"""
        try:
            config = self.requirements["tables"]["content"]
            
            logger.debug(f"📊 Форматирование содержимого таблицы...")
            
            # Применяем форматирование шрифта
            self._apply_font_formatting(paragraph, config)
            logger.debug(f"   ↳ Шрифт: {config['font_name']} {config['font_size']}pt")
            
            # Выравнивание по центру для содержимого таблиц
            paragraph.alignment = FormattingConstants.ALIGN_MAP[config["alignment"]]
            logger.debug(f"   ↳ Выравнивание: {config['alignment']}")
            
            # Настройки параграфа
            pf = paragraph.paragraph_format
            
            # Междустрочный интервал
            line_spacing = config.get("line_spacing", 1.0)
            if line_spacing in FormattingConstants.LINE_SPACING_MAP:
                pf.line_spacing_rule = FormattingConstants.LINE_SPACING_MAP[line_spacing]
                logger.debug(f"   ↳ Междустрочный интервал: {line_spacing}")
            
            # Убираем отступы для содержимого таблиц
            pf.first_line_indent = Cm(0)
            pf.left_indent = Cm(0)
            pf.right_indent = Cm(0)
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            
            logger.debug(f"✅ Содержимое таблицы отформатировано")

        except Exception as e:
            logger.error(f"❌ Ошибка форматирования содержимого таблицы: {e}")
            raise

    def format_table(self, table_element) -> None:
        """Форматирует саму таблицу согласно ГОСТ"""
        try:
            config = self.requirements["tables"]["table"]
            
            logger.info(f"📊 Форматирование таблицы...")
            
            # Выравнивание таблицы по центру
            if hasattr(table_element, 'alignment'):
                from docx.enum.table import WD_TABLE_ALIGNMENT
                if config["alignment"] == "center":
                    table_element.alignment = WD_TABLE_ALIGNMENT.CENTER
                    logger.debug(f"   ↳ Выравнивание таблицы: по центру")
                elif config["alignment"] == "left":
                    table_element.alignment = WD_TABLE_ALIGNMENT.LEFT
                    logger.debug(f"   ↳ Выравнивание таблицы: по левому краю")
            
            # Настройка ширины таблицы
            if config.get("width_auto", True):
                # Автоматическая ширина
                table_element.autofit = True
                logger.debug(f"   ↳ Автоматическая ширина таблицы")
            
            # Форматирование ячеек таблицы
            content_config = self.requirements["tables"]["content"]
            header_config = self.requirements["tables"]["header"]
            
            for i, row in enumerate(table_element.rows):
                is_header = (i == 0)  # Первая строка считается заголовком
                
                for cell in row.cells:
                    # Форматируем параграфы в ячейках
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():  # Только непустые параграфы
                            if is_header:
                                # Форматирование заголовка
                                self._apply_font_formatting(paragraph, header_config)
                                paragraph.alignment = FormattingConstants.ALIGN_MAP[header_config["alignment"]]
                            else:
                                # Форматирование обычного содержимого
                                self._apply_font_formatting(paragraph, content_config)
                                paragraph.alignment = FormattingConstants.ALIGN_MAP[content_config["alignment"]]
                            
                            # Убираем отступы в ячейках
                            pf = paragraph.paragraph_format
                            pf.first_line_indent = Cm(0)
                            pf.left_indent = Cm(0)
                            pf.right_indent = Cm(0)
                            pf.space_before = Pt(0)
                            pf.space_after = Pt(0)
                            
                            # Междустрочный интервал
                            line_spacing = content_config.get("line_spacing", 1.0)
                            if line_spacing in FormattingConstants.LINE_SPACING_MAP:
                                pf.line_spacing_rule = FormattingConstants.LINE_SPACING_MAP[line_spacing]
            
            logger.info(f"✅ Таблица отформатирована")

        except Exception as e:
            logger.error(f"❌ Ошибка форматирования таблицы: {e}")
            raise

    def format_figure_caption(self, paragraph) -> None:
        """Форматирует подпись рисунка согласно ГОСТ (без коррекции нумерации)"""
        try:
            config = self.requirements["figures"]["caption"]
            
            text = paragraph.text.strip()
            logger.info(f"🖼️ Форматирование подписи рисунка: {text}")
            
            # Применяем форматирование к существующему тексту (БЕЗ коррекции нумерации)
            self._apply_font_formatting(paragraph, config)
            
            # Выравнивание по центру (согласно ГОСТ)
            paragraph.alignment = FormattingConstants.ALIGN_MAP[config["alignment"]]
            logger.debug(f"   ↳ Выравнивание: {config['alignment']}")
            
            # Настройки параграфа
            pf = paragraph.paragraph_format
            
            # Отступы до и после подписи
            spacing_config = config["spacing"]
            pf.space_before = Pt(spacing_config["before_pt"])
            pf.space_after = Pt(spacing_config["after_pt"])
            logger.debug(f"   ↳ Отступы: до={spacing_config['before_pt']}pt, после={spacing_config['after_pt']}pt")
            
            # Убираем отступ первой строки для подписей
            pf.first_line_indent = Cm(0)
            pf.left_indent = Cm(0)
            pf.right_indent = Cm(0)
            
            # Междустрочный интервал
            line_spacing = config.get("line_spacing", 1.0)
            if line_spacing in FormattingConstants.LINE_SPACING_MAP:
                pf.line_spacing_rule = FormattingConstants.LINE_SPACING_MAP[line_spacing]
                logger.debug(f"   ↳ Междустрочный интервал: {line_spacing}")
            
            logger.info(f"✅ Подпись рисунка отформатирована: {paragraph.text[:50]}...")

        except Exception as e:
            logger.error(f"❌ Ошибка форматирования подписи рисунка: {e}")
            raise

    def format_figure_image(self, paragraph) -> None:
        """Форматирует изображение рисунка согласно ГОСТ"""
        try:
            config = self.requirements["figures"]["image"]
            
            logger.info(f"🖼️ Форматирование изображения рисунка...")
            
            # Выравнивание по центру (согласно ГОСТ)
            paragraph.alignment = FormattingConstants.ALIGN_MAP[config["alignment"]]
            logger.debug(f"   ↳ Выравнивание изображения: {config['alignment']}")
            
            # Настройки параграфа
            pf = paragraph.paragraph_format
            
            # Отступы до и после изображения
            spacing_config = config["spacing"]
            pf.space_before = Pt(spacing_config["before_pt"])
            pf.space_after = Pt(spacing_config["after_pt"])
            logger.debug(f"   ↳ Отступы изображения: до={spacing_config['before_pt']}pt, после={spacing_config['after_pt']}pt")
            
            # Убираем отступы для изображений
            pf.first_line_indent = Cm(0)
            pf.left_indent = Cm(0)
            pf.right_indent = Cm(0)
            
            logger.info(f"✅ Изображение рисунка отформатировано")

        except Exception as e:
            logger.error(f"❌ Ошибка форматирования изображения рисунка: {e}")
            raise

    def format_formula(self, paragraph) -> None:
        """Форматирует математическую формулу согласно ГОСТ"""
        try:
            config = self.requirements["formulas"]["formula"]
            
            text = paragraph.text.strip()
            logger.info(f"🔢 Форматирование формулы: {text[:60]}...")
            
            # Применяем форматирование шрифта
            self._apply_font_formatting(paragraph, config)
            logger.debug(f"   ↳ Шрифт: {config['font_name']} {config['font_size']}pt")
            
            # Выравнивание по центру (согласно ГОСТ)
            paragraph.alignment = FormattingConstants.ALIGN_MAP[config["alignment"]]
            logger.debug(f"   ↳ Выравнивание: {config['alignment']}")
            
            # Настройки параграфа
            pf = paragraph.paragraph_format
            
            # Отступы до и после формулы
            spacing_config = config["spacing"]
            pf.space_before = Pt(spacing_config["before_pt"])
            pf.space_after = Pt(spacing_config["after_pt"])
            logger.debug(f"   ↳ Отступы: до={spacing_config['before_pt']}pt, после={spacing_config['after_pt']}pt")
            
            # Убираем отступы первой строки для формул
            pf.first_line_indent = Cm(0)
            pf.left_indent = Cm(0)
            pf.right_indent = Cm(0)
            
            logger.info(f"✅ Формула отформатирована: {text[:40]}...")

        except Exception as e:
            logger.error(f"❌ Ошибка форматирования формулы: {e}")
            raise

    def format_formula_numbering(self, paragraph) -> None:
        """Форматирует нумерацию формулы согласно ГОСТ"""
        try:
            config = self.requirements["formulas"]["numbering"]
            
            text = paragraph.text.strip()
            logger.info(f"🔢 Форматирование нумерации формулы: {text}")
            
            # Применяем форматирование шрифта
            self._apply_font_formatting(paragraph, config)
            logger.debug(f"   ↳ Шрифт: {config['font_name']} {config['font_size']}pt")
            
            # Выравнивание по правому краю (согласно ГОСТ)
            paragraph.alignment = FormattingConstants.ALIGN_MAP[config["alignment"]]
            logger.debug(f"   ↳ Выравнивание: {config['alignment']}")
            
            # Настройки параграфа
            pf = paragraph.paragraph_format
            
            # Отступы для нумерации
            spacing_config = config["spacing"]
            pf.space_before = Pt(spacing_config["before_pt"])
            pf.space_after = Pt(spacing_config["after_pt"])
            logger.debug(f"   ↳ Отступы: до={spacing_config['before_pt']}pt, после={spacing_config['after_pt']}pt")
            
            # Убираем отступы первой строки
            pf.first_line_indent = Cm(0)
            pf.left_indent = Cm(0)
            pf.right_indent = Cm(0)
            
            logger.info(f"✅ Нумерация формулы отформатирована")

        except Exception as e:
            logger.error(f"❌ Ошибка форматирования нумерации формулы: {e}")
            raise

    def format_formula_explanation(self, paragraph) -> None:
        """Форматирует пояснения к переменным формулы согласно ГОСТ"""
        try:
            config = self.requirements["formulas"]["variables_explanation"]
            
            text = paragraph.text.strip()
            logger.info(f"🔤 Форматирование пояснения к формуле: {text[:60]}...")
            
            # Применяем форматирование шрифта
            self._apply_font_formatting(paragraph, config)
            logger.debug(f"   ↳ Шрифт: {config['font_name']} {config['font_size']}pt")
            
            # Выравнивание по левому краю с отступом (согласно ГОСТ)
            paragraph.alignment = FormattingConstants.ALIGN_MAP[config["alignment"]]
            logger.debug(f"   ↳ Выравнивание: {config['alignment']}")
            
            # Настройки параграфа
            pf = paragraph.paragraph_format
            
            # Отступ первой строки (красная строка для пояснений)
            pf.first_line_indent = Cm(config["indent_cm"])
            pf.left_indent = Cm(0)
            pf.right_indent = Cm(0)
            logger.debug(f"   ↳ Красная строка: {config['indent_cm']}см")
            
            # Междустрочный интервал
            line_spacing = config.get("line_spacing", 1.5)
            if line_spacing in FormattingConstants.LINE_SPACING_MAP:
                pf.line_spacing_rule = FormattingConstants.LINE_SPACING_MAP[line_spacing]
                logger.debug(f"   ↳ Междустрочный интервал: {line_spacing}")
            
            # Отступы до и после пояснений
            spacing_config = config["spacing"]
            pf.space_before = Pt(spacing_config["before_pt"])
            pf.space_after = Pt(spacing_config["after_pt"])
            logger.debug(f"   ↳ Отступы: до={spacing_config['before_pt']}pt, после={spacing_config['after_pt']}pt")
            
            logger.info(f"✅ Пояснение к формуле отформатировано")

        except Exception as e:
            logger.error(f"❌ Ошибка форматирования пояснения к формуле: {e}")
            raise
