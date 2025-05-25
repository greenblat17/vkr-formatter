from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
import re
import logging
from typing import Dict, Any, List, Tuple
from pathlib import Path
import colorlog
from abc import ABC, abstractmethod

# ============================================================================
# CONSTANTS AND CONFIGURATION
# ============================================================================

class FormattingConstants:
    """Константы для форматирования"""
    
    # Выравнивание текста
    ALIGN_MAP = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
    }
    
    # Междустрочные интервалы
    LINE_SPACING_MAP = {
        1.0: WD_LINE_SPACING.SINGLE,
        1.5: WD_LINE_SPACING.ONE_POINT_FIVE,
        2.0: WD_LINE_SPACING.DOUBLE
    }
    
    # Маркеры содержания
    CONTENT_HEADERS = [
        "СОДЕРЖАНИЕ", "ОГЛАВЛЕНИЕ", "CONTENTS", "TABLE OF CONTENTS"
    ]
    
    # Маркеры начала основного содержания
    MAIN_CONTENT_MARKERS = [
        "ВВЕДЕНИЕ", "ГЛАВА 1", "1. ВВЕДЕНИЕ", "1 ВВЕДЕНИЕ", 
        "CHAPTER 1", "РЕФЕРАТ", "ABSTRACT", "АННОТАЦИЯ"
    ]
    
    # Маркеры титульной страницы
    TITLE_PAGE_MARKERS = [
        "ДИПЛОМНАЯ РАБОТА", "ВЫПУСКНАЯ КВАЛИФИКАЦИОННАЯ РАБОТА", 
        "МИНИСТЕРСТВО ОБРАЗОВАНИЯ", "МИНИСТЕРСТВО НАУКИ",
        "ФЕДЕРАЛЬНОЕ ГОСУДАРСТВЕННОЕ", "ОБРАЗОВАТЕЛЬНОЕ УЧРЕЖДЕНИЕ",
        "ВЫСШЕГО ОБРАЗОВАНИЯ", "КАФЕДРА", "НАПРАВЛЕНИЕ ПОДГОТОВКИ",
        "ПРОФИЛЬ", "ТЕМА:", "ВЫПОЛНИЛ:", "СТУДЕНТ", "ГРУППЫ",
        "НАУЧНЫЙ РУКОВОДИТЕЛЬ", "КОНСУЛЬТАНТ", "ДОПУЩЕН К ЗАЩИТЕ",
        "РАБОТА ВЫПОЛНЕНА", "ОЦЕНКА", "ПОДПИСЬ"
    ]
    
    # Маркеры служебных разделов
    SERVICE_MARKERS = [
        "ЗАДАНИЕ НА", "КАЛЕНДАРНЫЙ ПЛАН", "КАЛЕНДАРНО-ТЕМАТИЧЕСКИЙ",
        "ТЕХНИЧЕСКОЕ ЗАДАНИЕ", "УТВЕРЖДАЮ", "РАССМОТРЕНО",
        "СОГЛАСОВАНО", "ОТЗЫВ", "РЕЦЕНЗИЯ", "СПРАВКА О ВНЕДРЕНИИ",
        "АКТ О ВНЕДРЕНИИ"
    ]

# ============================================================================
# LOGGING SETUP
# ============================================================================

def setup_colored_logging():
    """Настраивает цветное логирование"""
    color_formatter = colorlog.ColoredFormatter(
        '%(log_color)s%(asctime)s - %(levelname)-8s%(reset)s %(message)s',
        datefmt='%H:%M:%S',
        log_colors={
            'DEBUG': 'cyan', 'INFO': 'green', 'WARNING': 'yellow',
            'ERROR': 'red', 'CRITICAL': 'red,bg_white',
        },
        style='%'
    )
    
    handler = colorlog.StreamHandler()
    handler.setFormatter(color_formatter)
    
    logger = colorlog.getLogger(__name__)
    logger.setLevel(logging.INFO)
    logger.handlers.clear()
    logger.addHandler(handler)
    
    return logger

logger = setup_colored_logging()

# ============================================================================
# DOCUMENT STATE MANAGEMENT
# ============================================================================

class DocumentState:
    """Управляет состоянием обработки документа"""
    
    def __init__(self):
        self.in_title_section = True
        self.in_contents_section = False
        self.found_main_content = False
        self.pages_skipped = 0
    
    def start_contents_section(self):
        """Начинает раздел содержания"""
        self.in_contents_section = True
        logger.info("📑 Переход в режим содержания")
    
    def start_main_content(self):
        """Начинает основное содержание"""
        self.in_title_section = False
        self.in_contents_section = False
        self.found_main_content = True
        logger.info("🟢 Переход к основному содержанию")
    
    def is_in_service_section(self) -> bool:
        """Находимся ли в служебной секции"""
        return self.in_title_section or self.in_contents_section
    
    def get_state_info(self) -> Dict[str, bool]:
        """Возвращает информацию о текущем состоянии"""
        return {
            'in_title_section': self.in_title_section,
            'in_contents_section': self.in_contents_section,
            'found_main_content': self.found_main_content
        }

# ============================================================================
# CONTENT DETECTION AND ANALYSIS
# ============================================================================

class ContentDetector:
    """Определяет типы контента в документе"""
    
    @staticmethod
    def is_title_page_content(text: str) -> bool:
        """Определяет содержимое титульного листа"""
        text_upper = text.upper()
        
        # Проверка маркеров титульной страницы
        for marker in FormattingConstants.TITLE_PAGE_MARKERS:
            if marker in text_upper:
                return True
        
        # Паттерны для ФИО
        fio_patterns = [
            r"[А-ЯЁ][а-яё]+\s+[А-ЯЁ]\.[А-ЯЁ]\.",  # Иванов И.И.
            r"[А-ЯЁ][а-яё]+\s+[А-ЯЁ][а-яё]+\s+[А-ЯЁ][а-яё]+",  # Иванов Иван Иванович
        ]
        
        for pattern in fio_patterns:
            if re.search(pattern, text):
                return True
        
        # Короткие строки с высоким процентом заглавных букв
        if len(text) < 200:
            alpha_chars = [c for c in text if c.isalpha()]
            if alpha_chars:
                upper_ratio = sum(1 for c in alpha_chars if c.isupper()) / len(alpha_chars)
                if upper_ratio > 0.8 and len(text.split()) <= 5:
                    return True
        
        return False
    
    @staticmethod
    def is_contents_header(text: str) -> bool:
        """Определяет заголовок страницы содержания"""
        text_upper = text.upper().strip()
        return text_upper in FormattingConstants.CONTENT_HEADERS
    
    @staticmethod
    def is_contents_line(text: str) -> bool:
        """Определяет строку содержания с номерами страниц"""
        text_clean = text.strip()
        
        if not text_clean:
            return False
        
        # Паттерны для строк содержания
        content_patterns = [
            # С точками и номерами страниц
            r".+\.{3,}.+\d+$",  # "Введение...........3"
            r".+\.{2,}\s*\d+$",  # "1. Обзор литературы..5" 
            r"^[А-ЯЁ\d\.\s]+\.{3,}\d+$",  # "ГЛАВА 1...10"
            r"^\d+[\.\s][А-ЯЁа-яё\s]+\.{3,}\d+$",  # "1 Введение.....4"
            r"^\d+\.\d+[\.\s][А-ЯЁа-яё\s]+\.{3,}\d+$",  # "1.1 Подраздел.....8"
            
            # С пробелами и номерами страниц (без точек)
            r"^[А-ЯЁа-яё\s]+\s+\d+$",  # "Введение    8"
            r"^\d+\.\s*[А-ЯЁа-яё\s]+\s+\d+$",  # "1. Анализ предметной области   11"
            r"^\d+\s+[А-ЯЁа-яё\s]+\s+\d+$",  # "1 Введение   4"
            r"^\d+\.\d+\s+[А-ЯЁа-яё\s]+\s+\d+$",  # "1.1 Недостатки   11"
            
            # Специальные случаи
            r"^[А-ЯЁа-яё\s,]+\d+$",  # "Определения, обозначения и сокращения5"
        ]
        
        for pattern in content_patterns:
            if re.search(pattern, text_clean):
                return True
        
        # Дополнительная эвристика: строка заканчивается числом и содержит мало слов
        if re.search(r'\d+$', text_clean):
            words = text_clean.split()
            if 2 <= len(words) <= 8:
                try:
                    int(words[-1])
                    return True
                except ValueError:
                    pass
        
        return False
    
    @staticmethod
    def is_service_content(text: str) -> bool:
        """Определяет служебные разделы"""
        text_upper = text.upper()
        
        for marker in FormattingConstants.SERVICE_MARKERS:
            if marker in text_upper:
                return True
        
        return False
    
    @staticmethod
    def is_main_content_start(text: str) -> bool:
        """Определяет начало основного содержания"""
        text_upper = text.upper().strip()
        
        # Исключаем строки содержания
        if ContentDetector.is_contents_line(text):
            return False
        
        # Точные совпадения с маркерами
        for marker in FormattingConstants.MAIN_CONTENT_MARKERS:
            if text_upper == marker:
                return True
        
        # Паттерны для глав (без номеров страниц в конце)
        chapter_patterns = [
            r"^ГЛАВА\s+\d+$",  # "ГЛАВА 1"
            r"^\d+\.\s*[А-ЯЁ][А-ЯЁа-яё\s]*$",  # "1. ВВЕДЕНИЕ"
            r"^\d+\s+[А-ЯЁ][А-ЯЁа-яё\s]*$",    # "1 ВВЕДЕНИЕ"
        ]
        
        for pattern in chapter_patterns:
            if re.match(pattern, text_upper):
                if not re.search(r'\s+\d+$', text.strip()):
                    return True
        
        return False

# ============================================================================
# PARAGRAPH CLASSIFICATION
# ============================================================================

class ParagraphClassifier:
    """Классифицирует типы параграфов"""
    
    def __init__(self, requirements: Dict[str, Any]):
        self.requirements = requirements
        self.detector = ContentDetector()
        self.state = DocumentState()
    
    def classify_paragraph(self, text: str) -> str:
        """
        Классифицирует параграф
        
        Returns:
            str: "skip", "h1", "h2", "list", "regular"
        """
        if not text:
            return "skip"
        
        text_clean = text.strip()
        
        # 1. Проверяем заголовок содержания
        if self.detector.is_contents_header(text_clean):
            logger.info(f"📑 НАЧАЛО СОДЕРЖАНИЯ: {text_clean[:50]}...")
            self.state.start_contents_section()
            return "skip"
        
        # 2. Если в содержании
        if self.state.in_contents_section:
            return self._classify_in_contents_section(text_clean)
        
        # 3. Проверяем служебные разделы
        if (self.detector.is_title_page_content(text_clean) or 
            self.detector.is_service_content(text_clean)):
            logger.debug(f"🔴 СЛУЖЕБНЫЙ РАЗДЕЛ: {text_clean[:50]}...")
            return "skip"
        
        # 4. Проверяем начало основного содержания
        if (not self.state.found_main_content and 
            self.detector.is_main_content_start(text_clean)):
            logger.info(f"🟢 НАЧАЛО ОСНОВНОГО СОДЕРЖАНИЯ: {text_clean[:60]}...")
            self.state.start_main_content()
            return self._classify_content_paragraph(text_clean)
        
        # 5. Если в титульной секции
        if self.state.in_title_section:
            logger.debug(f"⚪ ПРОПУСК (титульная): {text_clean[:50]}...")
            return "skip"
        
        # 6. Классифицируем как обычное содержание
        return self._classify_content_paragraph(text_clean)
    
    def _classify_in_contents_section(self, text_clean: str) -> str:
        """Классифицирует параграфы в разделе содержания"""
        # Если это строка содержания - пропускаем
        if self.detector.is_contents_line(text_clean):
            logger.debug(f"📑 СОДЕРЖАНИЕ (строка): {text_clean[:50]}...")
            return "skip"
        
        # Если начинается основное содержание
        if self.detector.is_main_content_start(text_clean):
            logger.info(f"🟢 КОНЕЦ СОДЕРЖАНИЯ, НАЧАЛО ОСНОВНОГО СОДЕРЖАНИЯ: {text_clean[:60]}...")
            self.state.start_main_content()
            return self._classify_content_paragraph(text_clean)
        else:
            # Пустая строка или неопределенное в содержании
            logger.debug(f"📑 СОДЕРЖАНИЕ (прочее): {text_clean[:50]}...")
            return "skip"
    
    def _classify_content_paragraph(self, text_clean: str) -> str:
        """Классифицирует параграфы основного содержания"""
        if self._is_h1_paragraph(text_clean):
            return "h1"
        elif self._is_h2_paragraph(text_clean):
            return "h2"
        elif self._is_list_paragraph(text_clean):
            return "list"
        else:
            return "regular"
    
    def _is_h1_paragraph(self, text: str) -> bool:
        """Проверяет H1 заголовок"""
        patterns = self.requirements["h1_formatting"]["detection_patterns"]
        
        for pattern in patterns:
            if re.match(pattern, text.upper().strip()):
                return True
        
        # Дополнительная проверка: короткий текст с заглавными буквами
        if len(text) < 100:
            alpha_chars = [c for c in text if c.isalpha()]
            if alpha_chars:
                upper_ratio = sum(1 for c in alpha_chars if c.isupper()) / len(alpha_chars)
                if upper_ratio > 0.7:
                    return True
        
        return False
    
    def _is_h2_paragraph(self, text: str) -> bool:
        """Проверяет H2 заголовок"""
        patterns = self.requirements["h2_formatting"]["detection_patterns"]
        
        for pattern in patterns:
            if re.match(pattern, text.strip()):
                return True
        
        return False
    
    def _is_list_paragraph(self, text: str) -> bool:
        """Проверяет элемент списка"""
        patterns = self.requirements["lists"]["bullet_lists"]["detection_patterns"]
        
        for pattern in patterns:
            if re.match(pattern, text):
                return True
        
        return False
    
    def get_state(self) -> DocumentState:
        """Возвращает текущее состояние"""
        return self.state

# ============================================================================
# PARAGRAPH FORMATTING
# ============================================================================

class ParagraphFormatter:
    """Форматирует параграфы разных типов"""
    
    def __init__(self, requirements: Dict[str, Any]):
        self.requirements = requirements
    
    def format_h1(self, paragraph) -> None:
        """Форматирует заголовок H1"""
        try:
            config = self.requirements["h1_formatting"]
            
            # Разрыв страницы
            if config["page_break_before"] and self._not_first_paragraph(paragraph):
                self._add_page_break_before(paragraph)
            
            # Форматирование
            self._apply_font_formatting(paragraph, config)
            
            # Заглавные буквы
            if config["text_transform"] == "uppercase":
                self._make_text_uppercase(paragraph, config)
            
            # Выравнивание и отступы
            paragraph.alignment = FormattingConstants.ALIGN_MAP[config["alignment"]]
            
            pf = paragraph.paragraph_format
            pf.space_before = Pt(config["space_before_pt"])
            pf.space_after = Pt(config["space_after_pt"])
            
            logger.debug(f"H1 отформатирован: {paragraph.text[:30]}...")
            
        except Exception as e:
            logger.error(f"Ошибка форматирования H1: {e}")
            raise
    
    def format_h2(self, paragraph) -> None:
        """Форматирует заголовок H2"""
        try:
            config = self.requirements["h2_formatting"]
            
            self._apply_font_formatting(paragraph, config)
            paragraph.alignment = FormattingConstants.ALIGN_MAP[config["alignment"]]
            
            pf = paragraph.paragraph_format
            pf.space_before = Pt(config["space_before_pt"])
            pf.space_after = Pt(config["space_after_pt"])
            pf.left_indent = Cm(config.get("paragraph_indent_cm", 0))
            
            logger.debug(f"H2 отформатирован: {paragraph.text[:30]}...")
            
        except Exception as e:
            logger.error(f"Ошибка форматирования H2: {e}")
            raise
    
    def format_list(self, paragraph) -> None:
        """Форматирует элемент списка"""
        try:
            config = self.requirements["lists"]["bullet_lists"]
            font_config = config["font"]
            
            self._apply_font_formatting(paragraph, {
                "font_name": font_config["name"],
                "font_size": font_config["size"]
            })
            
            paragraph.alignment = FormattingConstants.ALIGN_MAP[config["alignment"]]
            
            pf = paragraph.paragraph_format
            pf.left_indent = Cm(config["indent_cm"])
            
            line_spacing = font_config["line_spacing"]
            if line_spacing in FormattingConstants.LINE_SPACING_MAP:
                pf.line_spacing_rule = FormattingConstants.LINE_SPACING_MAP[line_spacing]
            
            logger.debug(f"Список отформатирован: {paragraph.text[:30]}...")
            
        except Exception as e:
            logger.error(f"Ошибка форматирования списка: {e}")
            raise
    
    def format_regular(self, paragraph) -> None:
        """Форматирует обычный параграф"""
        try:
            if not paragraph.text.strip():
                return
            
            config = self.requirements["base_formatting"]
            
            self._apply_font_formatting(paragraph, config)
            paragraph.alignment = FormattingConstants.ALIGN_MAP[config["text_alignment"]]
            
            pf = paragraph.paragraph_format
            pf.first_line_indent = Cm(config["paragraph_indent_cm"])
            
            line_spacing = config["line_spacing"]
            if line_spacing in FormattingConstants.LINE_SPACING_MAP:
                pf.line_spacing_rule = FormattingConstants.LINE_SPACING_MAP[line_spacing]
            
        except Exception as e:
            logger.error(f"Ошибка форматирования обычного параграфа: {e}")
            raise
    
    def _apply_font_formatting(self, paragraph, config: Dict[str, Any]) -> None:
        """Применяет форматирование шрифта"""
        if not paragraph.runs:
            paragraph.add_run()
        
        for run in paragraph.runs:
            font = run.font
            
            if "font_name" in config:
                font.name = config["font_name"]
            
            if "font_size" in config:
                font.size = Pt(config["font_size"])
            
            if config.get("font_weight") == "bold":
                font.bold = True
    
    def _make_text_uppercase(self, paragraph, config: Dict[str, Any]) -> None:
        """Преобразует текст в верхний регистр"""
        original_text = paragraph.text
        paragraph.clear()
        run = paragraph.add_run(original_text.upper())
        
        font = run.font
        font.name = config["font_name"]
        font.size = Pt(config["font_size"])
        if config["font_weight"] == "bold":
            font.bold = True
    
    def _add_page_break_before(self, paragraph) -> None:
        """Добавляет разрыв страницы"""
        if paragraph.runs:
            first_run = paragraph.runs[0]
            first_run.add_break(WD_BREAK.PAGE)
        else:
            run = paragraph.add_run()
            run.add_break(WD_BREAK.PAGE)
    
    def _not_first_paragraph(self, target_paragraph) -> bool:
        """Проверяет, что параграф не первый"""
        try:
            doc = target_paragraph._parent
            while hasattr(doc, '_parent') and doc._parent is not None:
                doc = doc._parent
            
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph == target_paragraph:
                    for j in range(i):
                        if doc.paragraphs[j].text.strip():
                            return True
                    return False
            
            return True
            
        except Exception:
            return True

# ============================================================================
# STATISTICS TRACKING
# ============================================================================

class StatisticsTracker:
    """Отслеживает статистику обработки"""
    
    def __init__(self):
        self.stats = {
            'total_paragraphs': 0,
            'skipped_paragraphs': 0,
            'h1_formatted': 0,
            'h2_formatted': 0,
            'lists_formatted': 0,
            'regular_formatted': 0,
            'errors': 0
        }
    
    def increment(self, stat_name: str):
        """Увеличивает счетчик"""
        if stat_name in self.stats:
            self.stats[stat_name] += 1
    
    def get_statistics(self, state: DocumentState) -> Dict[str, Any]:
        """Возвращает полную статистику"""
        stats = self.stats.copy()
        stats.update({
            'title_pages_detected': 1 if state.found_main_content else 0,
            'main_content_found': state.found_main_content,
            'contents_section_detected': not state.in_contents_section and state.found_main_content
        })
        return stats

# ============================================================================
# MAIN FORMATTER CLASS
# ============================================================================

class VKRFormatter:
    """Основной класс для форматирования ВКР"""
    
    def __init__(self, requirements: Dict[str, Any]):
        self.requirements = requirements
        self.classifier = ParagraphClassifier(requirements)
        self.formatter = ParagraphFormatter(requirements)
        self.stats = StatisticsTracker()
    
    def format_document(self, input_path: str, output_path: str) -> bool:
        """Форматирует документ"""
        try:
            logger.info(f"📂 Начинаем форматирование: {input_path}")
            logger.info(f"💾 Выходной путь: {output_path}")
            
            # Проверяем входной файл
            input_file = Path(input_path)
            if not input_file.exists():
                logger.error(f"❌ Входной файл не существует: {input_path}")
                return False
            
            # Загружаем документ
            logger.info("📖 Загружаем документ...")
            doc = Document(input_path)
            logger.info(f"✅ Документ загружен, параграфов: {len(doc.paragraphs)}")
            
            # Применяем глобальные настройки
            logger.info("⚙️  Применяем глобальные настройки...")
            self._apply_global_settings(doc)
            
            # Обрабатываем параграфы
            logger.info("🔄 Обрабатываем параграфы...")
            self._process_all_paragraphs(doc)
            
            # Сохраняем результат
            logger.info(f"💾 Сохраняем документ в: {output_path}")
            doc.save(output_path)
            
            # Проверяем результат
            output_file = Path(output_path)
            if output_file.exists():
                logger.info(f"✅ Файл создан, размер: {output_file.stat().st_size} байт")
            else:
                logger.error(f"❌ Файл НЕ создался: {output_path}")
                return False
            
            final_stats = self.get_statistics()
            logger.info(f"🎉 Форматирование завершено! Статистика: {final_stats}")
            return True
            
        except Exception as e:
            logger.error(f"Ошибка форматирования: {e}")
            import traceback
            logger.error(f"Полная трассировка: {traceback.format_exc()}")
            return False
    
    def _apply_global_settings(self, doc: Document) -> None:
        """Применяет глобальные настройки документа"""
        try:
            margins = self.requirements["base_formatting"]["margins_cm"]
            
            for section in doc.sections:
                section.top_margin = Cm(margins["top"])
                section.bottom_margin = Cm(margins["bottom"])
                section.left_margin = Cm(margins["left"])
                section.right_margin = Cm(margins["right"])
            
            logger.info(f"Применены поля: {margins}")
            
        except Exception as e:
            logger.error(f"Ошибка применения глобальных настроек: {e}")
            self.stats.increment('errors')
    
    def _process_all_paragraphs(self, doc: Document) -> None:
        """Обрабатывает все параграфы документа"""
        logger.info("Начинаем обработку параграфов...")
        
        for i, paragraph in enumerate(doc.paragraphs):
            self.stats.increment('total_paragraphs')
            
            try:
                text = paragraph.text.strip()
                paragraph_type = self.classifier.classify_paragraph(text)
                
                # Логируем непустые параграфы
                if text:
                    logger.debug(f"Параграф {i+1}: тип='{paragraph_type}', текст='{text[:100]}{'...' if len(text) > 100 else ''}'")
                
                # Применяем форматирование
                self._apply_paragraph_formatting(paragraph, paragraph_type, i+1, text)
                
            except Exception as e:
                logger.warning(f"Ошибка обработки параграфа {i+1}: {e}")
                self.stats.increment('errors')
        
        final_stats = self.stats.stats
        logger.info(f"Обработка параграфов завершена. Статистика: {final_stats}")
    
    def _apply_paragraph_formatting(self, paragraph, paragraph_type: str, index: int, text: str) -> None:
        """Применяет форматирование к параграфу"""
        if paragraph_type == "skip":
            self.stats.increment('skipped_paragraphs')
            logger.info(f"⏭️  ПРОПУСК #{index}: {text[:60]}{'...' if len(text) > 60 else ''}")
            
        elif paragraph_type == "h1":
            self.formatter.format_h1(paragraph)
            self.stats.increment('h1_formatted')
            logger.info(f"📝 H1 #{index}: {text[:40]}...")
            
        elif paragraph_type == "h2":
            self.formatter.format_h2(paragraph)
            self.stats.increment('h2_formatted')
            logger.info(f"📄 H2 #{index}: {text[:40]}...")
            
        elif paragraph_type == "list":
            self.formatter.format_list(paragraph)
            self.stats.increment('lists_formatted')
            logger.debug(f"📋 СПИСОК #{index}: {text[:40]}...")
            
        else:  # regular
            self.formatter.format_regular(paragraph)
            self.stats.increment('regular_formatted')
    
    def get_statistics(self) -> Dict[str, Any]:
        """Возвращает статистику обработки"""
        return self.stats.get_statistics(self.classifier.get_state())

# ============================================================================
# MAIN API FUNCTION
# ============================================================================

def format_vkr_document(input_path: str, requirements: Dict[str, Any], output_path: str) -> Tuple[bool, Dict[str, Any]]:
    """
    Форматирует ВКР согласно требованиям
    
    Args:
        input_path: путь к исходному файлу ВКР
        requirements: словарь требований
        output_path: путь к результирующему файлу
        
    Returns:
        tuple: (успех, статистика)
    """
    formatter = VKRFormatter(requirements)
    success = formatter.format_document(input_path, output_path)
    stats = formatter.get_statistics()
    
    return success, stats 