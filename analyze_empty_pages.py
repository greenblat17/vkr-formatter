#!/usr/bin/env python3
"""
Анализ пустых страниц в документе после добавления разрывов страниц
"""

from docx import Document
from docx.enum.text import WD_BREAK
import tempfile
import os

from vkr_formatter import format_vkr_document
from requirements_stub import analyze_requirements_stub


def create_test_document():
    """Создает тестовый документ для анализа пустых страниц"""
    doc = Document()
    
    # Титульная страница
    doc.add_paragraph("МИНИСТЕРСТВО ОБРАЗОВАНИЯ И НАУКИ РОССИЙСКОЙ ФЕДЕРАЦИИ")
    doc.add_paragraph("ВЫПУСКНАЯ КВАЛИФИКАЦИОННАЯ РАБОТА")
    doc.add_paragraph("Тема: Разработка системы автоматизации")
    
    # Содержание
    doc.add_page_break()
    doc.add_paragraph("СОДЕРЖАНИЕ")
    doc.add_paragraph("1. Введение    5")
    doc.add_paragraph("2. Проектирование системы    10")
    
    # Основное содержание
    doc.add_page_break()
    
    # H1 заголовки
    h1_1 = doc.add_paragraph("1. ВВЕДЕНИЕ")
    h1_1.style = doc.styles['Heading 1']
    doc.add_paragraph("Текст введения.")
    
    h1_2 = doc.add_paragraph("2. ПРОЕКТИРОВАНИЕ СИСТЕМЫ")
    h1_2.style = doc.styles['Heading 1']
    doc.add_paragraph("Текст основной части.")
    
    h1_3 = doc.add_paragraph("3. ЗАКЛЮЧЕНИЕ")
    h1_3.style = doc.styles['Heading 1']
    doc.add_paragraph("Текст заключения.")
    
    return doc


def analyze_document_structure(doc_path, title="АНАЛИЗ СТРУКТУРЫ ДОКУМЕНТА"):
    """Анализирует структуру документа и выявляет пустые страницы"""
    doc = Document(doc_path)
    
    print(f"\n📊 {title}")
    print("=" * 70)
    print(f"   Всего параграфов: {len(doc.paragraphs)}")
    
    print(f"\n{'№':<3} {'Текст':<45} {'Стиль':<15} {'Разрыв':<8} {'Пустой?':<8}")
    print("-" * 80)
    
    empty_pages_detected = []
    current_page_content = []
    page_number = 1
    
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        style_name = paragraph.style.name if paragraph.style else "None"
        has_page_break = False
        
        # Проверяем наличие разрыва страницы
        if paragraph.runs:
            for run in paragraph.runs:
                for element in run._element:
                    if 'br' in element.tag:
                        break_type = element.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type')
                        if break_type == 'page':
                            has_page_break = True
                            break
        
        # Если есть разрыв страницы, анализируем предыдущую страницу
        if has_page_break:
            if not any(content.strip() for content in current_page_content):
                empty_pages_detected.append(page_number)
            current_page_content = []
            page_number += 1
        
        # Добавляем содержимое текущего параграфа
        current_page_content.append(text)
        
        # Определяем, пустой ли параграф
        is_empty = not text
        
        break_indicator = "✅" if has_page_break else "❌"
        empty_indicator = "⚠️" if is_empty else "✅"
        
        print(f"{i+1:2d}. {text[:45]:<45} {style_name:<15} {break_indicator:<8} {empty_indicator:<8}")
    
    # Анализируем последнюю страницу
    if not any(content.strip() for content in current_page_content):
        empty_pages_detected.append(page_number)
    
    if empty_pages_detected:
        print(f"\n⚠️  ОБНАРУЖЕНЫ ПОТЕНЦИАЛЬНО ПУСТЫЕ СТРАНИЦЫ: {empty_pages_detected}")
    else:
        print(f"\n✅ ПУСТЫХ СТРАНИЦ НЕ ОБНАРУЖЕНО")
    
    return empty_pages_detected


def test_empty_pages_issue():
    """Тестирует проблему с пустыми страницами"""
    print("🔍 АНАЛИЗ ПРОБЛЕМЫ С ПУСТЫМИ СТРАНИЦАМИ")
    print("=" * 70)
    
    with tempfile.TemporaryDirectory() as tmpdir:
        # Создаем тестовый документ
        print("📝 Создаем тестовый документ...")
        demo_doc = create_test_document()
        input_path = os.path.join(tmpdir, "empty_pages_test_input.docx")
        demo_doc.save(input_path)
        
        # Анализируем исходный документ
        analyze_document_structure(input_path, "ИСХОДНЫЙ ДОКУМЕНТ")
        
        requirements = analyze_requirements_stub("test")
        
        # Форматируем документ
        print("\n🔄 Форматируем документ...")
        output_path = os.path.join(tmpdir, "empty_pages_test_output.docx")
        success, stats = format_vkr_document(
            input_path, requirements, output_path, 
            use_style_based=True, strict_style_mode=True
        )
        
        if success:
            # Сохраняем результат
            import shutil
            shutil.copy2(output_path, "empty_pages_analysis_result.docx")
            print(f"   📋 Результат сохранен: empty_pages_analysis_result.docx")
            
            # Анализируем результат
            empty_pages = analyze_document_structure(output_path, "ОТФОРМАТИРОВАННЫЙ ДОКУМЕНТ")
            
            print(f"\n📊 РЕЗУЛЬТАТЫ АНАЛИЗА:")
            print(f"   H1 заголовков отформатировано: {stats.get('h1_formatted', 0)}")
            print(f"   Пустых страниц обнаружено: {len(empty_pages)}")
            
            if empty_pages:
                print(f"\n⚠️  ПРОБЛЕМА: Обнаружены пустые страницы!")
                print(f"   Номера страниц: {empty_pages}")
                print(f"\n💡 ВОЗМОЖНЫЕ ПРИЧИНЫ:")
                print(f"   1. Разрыв страницы добавляется в неправильное место")
                print(f"   2. Разрыв создает пустую страницу перед заголовком")
                print(f"   3. Неправильная обработка пустых параграфов")
                return False
            else:
                print(f"\n✅ ОТЛИЧНО: Пустых страниц не обнаружено!")
                return True
        else:
            print("❌ Ошибка форматирования")
            return False


def analyze_page_break_placement():
    """Анализирует размещение разрывов страниц"""
    print("\n🔍 АНАЛИЗ РАЗМЕЩЕНИЯ РАЗРЫВОВ СТРАНИЦ")
    print("=" * 70)
    
    # Создаем простой документ для анализа
    doc = Document()
    
    # Тест 1: Разрыв в начале run
    p1 = doc.add_paragraph()
    run1 = p1.add_run()
    run1.add_break(WD_BREAK.PAGE)
    run1.add_text("Заголовок после разрыва")
    
    # Тест 2: Разрыв в отдельном run
    p2 = doc.add_paragraph()
    break_run = p2.add_run()
    break_run.add_break(WD_BREAK.PAGE)
    text_run = p2.add_run("Заголовок в отдельном run")
    
    # Тест 3: Пустой параграф с разрывом
    p3 = doc.add_paragraph()
    p3.add_run().add_break(WD_BREAK.PAGE)
    p4 = doc.add_paragraph("Заголовок после пустого параграфа")
    
    doc.save("page_break_placement_test.docx")
    
    print("📋 Создан файл: page_break_placement_test.docx")
    print("📊 Структура тестового документа:")
    
    for i, paragraph in enumerate(doc.paragraphs):
        print(f"   Параграф {i+1}: '{paragraph.text}'")
        if paragraph.runs:
            for j, run in enumerate(paragraph.runs):
                has_break = any('br' in elem.tag for elem in run._element)
                print(f"      Run {j+1}: '{run.text}' (разрыв: {has_break})")


def main():
    """Главная функция"""
    print("=" * 80)
    print("🔍 АНАЛИЗ ПРОБЛЕМЫ С ПУСТЫМИ СТРАНИЦАМИ")
    print("=" * 80)
    
    try:
        success = test_empty_pages_issue()
        analyze_page_break_placement()
        
        print("\n" + "=" * 80)
        if not success:
            print("⚠️  ПРОБЛЕМА С ПУСТЫМИ СТРАНИЦАМИ ПОДТВЕРЖДЕНА")
            print("🔧 ТРЕБУЕТСЯ ИСПРАВЛЕНИЕ ЛОГИКИ РАЗРЫВОВ СТРАНИЦ")
            print("\n💡 ВОЗМОЖНЫЕ РЕШЕНИЯ:")
            print("   1. Изменить способ добавления разрыва страницы")
            print("   2. Добавить постобработку для удаления пустых страниц")
            print("   3. Использовать свойства параграфа вместо разрывов")
        else:
            print("✅ ПУСТЫХ СТРАНИЦ НЕ ОБНАРУЖЕНО")
        print("=" * 80)
        
    except Exception as e:
        print(f"💥 Ошибка анализа: {e}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main() 